# -*- coding: utf-8 -*-

"""
Module implementing CLASS.
导入模块、界面
"""
from PyQt5.QtGui import QTextBlockFormat, QTextCursor,QDesktopServices   #引入段落默认格式模块
from PyQt5.QtCore import pyqtSlot, Qt, QTimer,QDate, pyqtSignal, QObject, QDir, QFileInfo, QTranslator,QUrl
from PyQt5.QtWidgets import QMessageBox, QMainWindow, QApplication, QListWidgetItem,  QTableWidgetItem, QFileDialog, QDialog
from Ui_mainWindow import Ui_mainWindow
from sys import argv, exit
from pickle import load, dump
from docxtpl import DocxTemplate, Listing
from os import getcwd, path, makedirs
from Ui_Dialog import Ui_Dialog
from datetime import datetime
import docx, xlrd, json, re,requests

VERSION = "2.7.0"  # 本地版本号


def check_for_updates(current_version: str):
    """检查是否有新版本"""
    try:
        # 从 GitHub API 获取最新版本信息
        url = "https://api.github.com/repos/kevinyoung-zxc/fluffy-potato/releases/latest"
        response = requests.get(url)
        if response.status_code == 200:
            latest_release = response.json()
            latest_version = latest_release['tag_name']  # 例如 "v1.0.0"
            latest_version = latest_version.lstrip('v')  # 去掉 "v" 前缀
            
            # 比较版本号
            if latest_version > current_version:
                # 提示用户更新
                reply = QMessageBox.question(
                    None, '更新提示',
                    f'发现新版本 {latest_version}，是否前往下载？',
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.Yes
                )
                if reply == QMessageBox.Yes:
                    import webbrowser
                    webbrowser.open(latest_release['html_url'])  # 打开下载页面
            else:
                QMessageBox.information(None, '更新检查', '当前已是最新版本')
        else:
            QMessageBox.warning(None, '错误', '无法检查更新，请检查网络连接')
    except Exception as e:
        QMessageBox.critical(None, '错误', f'更新检查失败: {str(e)}')


def Accus_explode(Accuname,Accus):#案由解读函数，从案由中解读出违法条款和处罚依据
    if Accuname in Accus :
        GT_Accus={key:value for key, value in Accus.items()if key == Accuname}
        DicGT_Accus=dict(GT_Accus[Accuname])#空值出错
        list_Law=DicGT_Accus['违反条款'].split('|')
        list_Law.pop()
        list_Source=DicGT_Accus['依据条款'].split('|')
        list_Source.pop()
    else:
        list_Law=[]
        list_Source=[]
    return (list_Law, list_Source)
def Accus_join(Accus_list):
    fawdict={}
    for i in Accus_list[0::2]:
        Ilist=i.split('》')
        Ilist[0]+='》'
        if Ilist[0] in fawdict.keys():
            value=fawdict[Ilist[0]]
            fawdict[Ilist[0]]=value+'、'+Ilist[1]
        else:
            fawdict[Ilist[0]]=Ilist[1]
    listItme=[]
    for key in fawdict.keys():
        key+=fawdict[key]
        listItme.append(key)
    Caselawitem='&#&@&'.join(listItme)
    Caselawtext=''
    if len(Accus_list[0::2]) ==1:
        Caselawitem=Caselawitem.replace('&#&@&', '')
    elif len(Accus_list[0::2]) >2:
        Caselawitem=Caselawitem.replace('&#&@&', '、', len(listItme)-2)
        Caselawitem=Caselawitem.replace('&#&@&', '和')
    elif len(Accus_list[0::2]) ==2:
        Caselawitem=Caselawitem.replace('&#&@&', '和')
    for (j, k) in zip(Accus_list[0::2], Accus_list[1::2]):
        Caselawtext += j+'：'+k+'\a'
    Caselawtext=Caselawtext.strip('\a')
    return(Caselawitem, Caselawtext)

def open_folder(path_output):
    url = QUrl.fromLocalFile(path_output)
    QDesktopServices.openUrl(url)
    
def compare_dates(input_date_str, default_date_str):
    # 将输入的日期字符串和默认日期字符串转换为 datetime.date 对象
    input_date = datetime.strptime(input_date_str, "%Y-%m-%d").date()
    default_date = datetime.strptime(default_date_str, "%Y-%m-%d").date()

    # 比较两个日期，返回较早的日期字符串
    if input_date <= default_date:
        return input_date_str
    else:
        return default_date_str    
    
##定义信号类——————————————————
class Signal(QObject):
    bridge = pyqtSignal()
    context_save = pyqtSignal()
##定义信号管理类负责信号对象的创建和使用分离---------------------------
class SignalManager:
    def __init__(self):
        self.signal = Signal()  # 创建 Signal 实例



####----主窗口————————————————————————————————————
class CLASS(QMainWindow, Ui_mainWindow):
    """
    Class documentation goes here.
    """
    def __init__(self, signal_manager,parent=None, ):
        """
        Constructor
        @param parent reference to the parent widget
        @type QWidget
        """
        super(CLASS, self).__init__(parent)
        self.setupUi(self)

        self.cursor_shift = True  # 光标切换开关
        self.block_shift_false = 0
        self.block_shift_true = 0
        #-------表格插件最后一列自动扩展-----
        self.tableWidget_laws.horizontalHeader().setStretchLastSection(True)
        self.tableWidget_lawsources.horizontalHeader().setStretchLastSection(True)
        #初始化制作提示信息变量
        self.MadeNote =''
        #-初始化证据材料、证据证明内容、违法事实等组成list------------------------------#
        
        self.TextEdit_Evidence_list=[self.textEdit_evidencelist_0,self.textEdit_evidencelist_1,self.textEdit_evidencelist_2,\
                                self.textEdit_evidencelist_3,self.textEdit_evidencelist_4,self.textEdit_evidencelist_5,\
                                    self.textEdit_evidencelist_6,self.textEdit_evidencelist_7]
        self.TextEdit_Evidenceproof_list=[self.textEdit_evidenceproof_0,self.textEdit_evidenceproof_1,self.textEdit_evidenceproof_2,\
                                     self.textEdit_evidenceproof_3,self.textEdit_evidenceproof_4,self.textEdit_evidenceproof_5,\
                                         self.textEdit_evidenceproof_6,self.textEdit_evidenceproof_7]
        self.TextEdit_EvidenceGZ_list=[self.textEdit_evidencelistGZ_0,self.textEdit_evidencelistGZ_1,self.textEdit_evidencelistGZ_2,\
                                self.textEdit_evidencelistGZ_3,self.textEdit_evidencelistGZ_4,self.textEdit_evidencelistGZ_5,\
                                    self.textEdit_evidencelistGZ_6,self.textEdit_evidencelistGZ_7]
        self.TextEdit_EvidenceproofGZ_list=[self.textEdit_evidenceproofGZ_0,self.textEdit_evidenceproofGZ_1,self.textEdit_evidenceproofGZ_2,\
                                     self.textEdit_evidenceproofGZ_3,self.textEdit_evidenceproofGZ_4,self.textEdit_evidenceproofGZ_5,\
                                         self.textEdit_evidenceproofGZ_6,self.textEdit_evidenceproofGZ_7]
        self.TextEdit_Illfact_listA=[self.textEdit_illfacts_A0,self.textEdit_illfacts_A1,self.textEdit_illfacts_A2,\
                                self.textEdit_illfacts_A3,self.textEdit_illfacts_A4]
        self.TextEdit_Illfact_listB=[self.textEdit_illfacts_B0,self.textEdit_illfacts_B1,self.textEdit_illfacts_B2,\
                                self.textEdit_illfacts_B3,self.textEdit_illfacts_B4]
        self.TextEdit_Illfact_listC=[self.textEdit_illfacts_C0,self.textEdit_illfacts_C1,self.textEdit_illfacts_C2,\
                                self.textEdit_illfacts_C3,self.textEdit_illfacts_C4]
        self.TextEdit_Illfact_listD=[self.textEdit_illfacts_D0,self.textEdit_illfacts_D1,self.textEdit_illfacts_D2,\
                                self.textEdit_illfacts_D3,self.textEdit_illfacts_D4]
        self.TextEdit_Illfact_listE=[self.textEdit_illfacts_E0,self.textEdit_illfacts_E1,self.textEdit_illfacts_E2,\
                                self.textEdit_illfacts_E3,self.textEdit_illfacts_E4]
        self.TextEdit_Illfact_listGZA=[self.textEdit_illfactsGZ_A0,self.textEdit_illfactsGZ_A1,self.textEdit_illfactsGZ_A2]
        self.TextEdit_Illfact_listGZB=[self.textEdit_illfactsGZ_B0,self.textEdit_illfactsGZ_B1,self.textEdit_illfactsGZ_B2]
        self.TextEdit_Illfact_listGZC=[self.textEdit_illfactsGZ_C0,self.textEdit_illfactsGZ_C1,self.textEdit_illfactsGZ_C2]
        self.TextEdit_Illfact_listGZD=[self.textEdit_illfactsGZ_D0,self.textEdit_illfactsGZ_D1,self.textEdit_illfactsGZ_D2]
        self.TextEdit_Illfact_listGZE=[self.textEdit_illfactsGZ_E0,self.textEdit_illfactsGZ_E1,self.textEdit_illfactsGZ_E2]
        #------------------------------------------------------------------------------------
        self.Accus = [] #案由列表函数，管理案由
        self.Accus_initA()#加载案由列表
        #加载历史记录-----------------------
        self.History = {}  #读取信息
        self.context={}
        self.AccuLists=[]
        self.Case_illeglFacts = ''
        try:
            with open ('History.pickle','rb') as History_file:
                self.History = load(History_file)
        except:
            self.History = {'Fst_History':'','Sed_History':'', 'Tre_History':'' , 'DeSetone':'' , 'DeSettwo':''}
        #加载缓存文件
        try:
            self.READ_context('AutoSave.pickle')#读取缓存文件
        except:
            self.dateEdit_execuTime.setDate(QDate.currentDate())
            self.dateEdit_GZtime.setDate(QDate.currentDate())
            self.dateEdit_filingtimeR.setDate(QDate.currentDate())
            self.dateEdit_filingtime.setDate(QDate.currentDate())
            self.dateEdit_JDTime.setDate(QDate.currentDate())
        #在 CLASS 类中创建 ChildWindow 实例
        #定义信号
        self.signal_manager = signal_manager  # 依赖注入 SignalManager
        self.signal_manager.signal.bridge.connect(lambda:self.READ_context('AutoSave.pickle'))
        self.signal_manager.signal.context_save.connect(lambda:self.SAVE_context('AutoSave.pickle'))
         # 在 CLASS 中创建 ChildWindow 实例
        self.child_window = ChildWindow(self.signal_manager, self)
        # 绑定按钮点击事件
        self.pushButton_replace.clicked.connect(self.show_child_window)
        # 初始化光标切换和段落格式控制变量,用于解决撤销失效问题------------------------------#
        
        #加载主体信息数据表格---------------
        if path.exists('./Date/主体信息/') == False :
            makedirs('./Date/主体信息/')
        xlsdir = QDir('./Date/主体信息/')
        xlslist = xlsdir.entryInfoList()
        xlsnamelist=[]
        for i in xlslist:
            i = QFileInfo.fileName(i)
            if ".xls" in i :
                xlsnamelist.append(i)
        if xlsnamelist:
            self.comboBox_Type.clear()
            self.comboBox_Type.addItems(xlsnamelist)
            self.comboBox_Type.setCurrentIndex(-1)
        self.show()

        #---------自动保存模块----------#
        Auto_time=QTimer(self)
        try:
            Auto_time.timeout.connect(lambda:self.SAVE_context('AutoSave.pickle'))
        except:
            QMessageBox.information(self,'错误提示', '自动保存失败，请核对案由格式和内容。')
            Auto_time.stop()
        Auto_time.start(30000)
        
    def show_child_window(self):
        """显示子窗口"""
        self.child_window.show()
        self.child_window.raise_()
        self.child_window.showNormal()

    def WORD_tpl(self, DOTA,OUTPUT,TECH, SUFFIX):#赋值函数
        """将界面文本值赋予模版文件"""
        NowDate=QDate.currentDate()
        NowDate=NowDate.toString(Qt.DefaultLocaleLongDate)
        OUTPUT_suffix=OUTPUT+'_'+NowDate
        doc=DocxTemplate(r"Date\%s.docx"%DOTA)
        doctext={ key:value for key,value in self.context.items() if key in TECH}
        doc.render(doctext)
        try:
            data='./Output/'+SUFFIX
            print(data, path.exists(data))
            if path.exists(data) == False :
                makedirs(data)
            doc.save(r"Output/%s/%s.docx"%(SUFFIX, OUTPUT_suffix))
            self.MadeNote +='%s文书制作成功。\n'%OUTPUT
        except :
            self.MadeNote +='%s文书制作失败。\n'%OUTPUT
            QMessageBox.information(self,'错误提示', '请先关闭已打开的WORD文档')
    #段落格式函数
    def blockFormat_initA(self, list_textEditName,Initialize_Format):
        """
        实现敲回车（段落数增加）刷新段落格式
        Slot documentation goes here.
        list_textEditName 为textEdit名字列表       
        """
        Format =QTextBlockFormat()
        Format.setLineHeight(150, 1) #行高1.5倍
        Format.setTextIndent(30)#首行缩进30
        for i in list_textEditName:
            doc= i.document()
            maxnum = doc.blockCount()  
            if Initialize_Format == False :
                if self.cursor_shift == True:
                    self.block_shift_true = int(maxnum)
                    if self.block_shift_true > self.block_shift_false: 
                        for i in range (0, maxnum):
                            cursor = QTextCursor(doc.findBlockByNumber(i))
                            cursor.setBlockFormat(Format)
                        self.block_shift_false = self.block_shift_true
                    self.cursor_shift = False
                else :
                    self.block_shift_false =int(maxnum)
                    if self.block_shift_true < self.block_shift_false: 
                        for i in range (0, maxnum):
                            cursor = QTextCursor(doc.findBlockByNumber(i))
                            cursor.setBlockFormat(Format)
                        self.block_shift_true = self.block_shift_false
                    self.cursor_shift = True
            elif Initialize_Format :
                 for i in range (0, maxnum):
                    cursor = QTextCursor(doc.findBlockByNumber(i))
                    cursor.setBlockFormat(Format)
                    
    #历史输入刷新函数
    def History_initA(self, History_key):
        FT_History={key:value for key, value in self.History.items()if key == 'Fst_History'}
        ST_History={key:value for key, value in self.History.items()if key == 'Sed_History'}
        TT_History={key:value for key, value in self.History.items()if key == 'Tre_History'}
        DicFst_History=dict(FT_History['Fst_History'])#空值出错
        DicSed_History=dict(ST_History['Sed_History'])
        DicTre_History=dict(TT_History['Tre_History'])
        self.textEdit_history1.setText('')#清空原值
        self.textEdit_history2.setText('')
        self.textEdit_history3.setText('')
        try:
            self.textEdit_history1.setText(DicFst_History[History_key])
            self.textEdit_history2.setText(DicSed_History[History_key])
            self.textEdit_history3.setText(DicTre_History[History_key])
        except:
            pass#print(f"{History_key}无历史输入记录。")
    #案由刷新函数
    def Accus_initA(self):
        try :
            with open ('Accus.pickle', 'rb') as Accus_file :
                self.Accus = load(Accus_file)
                self.listWidget_AccuEdit.clear()
                for i in self.Accus:
                    self.listWidget_AccuEdit.addItem(i)
        except :
            QMessageBox.information(self,'错误提示', '案由载入失败，请先通过案由设置新增案由。')
    #读取数据函数#存储对象赋值给界面
    def READ_context(self,File_name):
        with open (File_name,'rb') as context_file:
            self.context=load(context_file)
            self.AccuLists=self.context['Face_Accuname'].split('；')
            #案由载入检测
            AccuTest = True
            if self.AccuLists !=['']:
                for i in self.AccuLists:
                    if i in self.Accus:
                        pass
                    else :
                        QMessageBox.information(self,'错误提示', f"该案案由“{i}”未在软件案由列表内，请导入该案由后重试")
                        AccuTest = False
                        break
            if AccuTest :
                self.textEdit_litigant.setText(self.context['Case_name'])
                self.textEdit_accuss.setText(self.context['Face_Accuname'])
                self.comboBox_accusGZ.clear()#清空原值
                self.comboBox_accusDC.clear()#清空原值
                self.comboBox_accusDC.addItems(self.AccuLists)
                self.comboBox_accusGZ.addItems(self.AccuLists)
                self.comboBox_source.setCurrentIndex(self.context['Face_SourceIndex'])
                self.comboBox_accusDC.setCurrentIndex(self.context['Face_AccuDCIndex'])
                self.comboBox_accusGZ.setCurrentIndex(self.context['Face_AccuGZIndex'])
                self.dateEdit_filingtime.setDate(self.context['Face_time'])
                self.textEdit_absractcase.setText(self.context['Case_abstract'])
                self.textEdit_process.setText(self.context['Face_process'])
                #修改#----------赋予证据组表格数据------
                for i in range(8):
                    try:
                        Evidence = self.TextEdit_Evidence_list[i]
                        Evidence.setText(self.context['Face_evidence'+str(i)])
                        Evidenceproof = self.TextEdit_Evidenceproof_list[i]
                        Evidenceproof.setText(self.context['Face_evidencedata'+str(i)])
                        EvidenceGZ = self.TextEdit_EvidenceGZ_list[i]
                        EvidenceGZ.setText(self.context['Face_evidenceGZ'+str(i)])
                        EvidenceproofGZ = self.TextEdit_EvidenceproofGZ_list[i]
                        EvidenceproofGZ.setText(self.context['Face_evidencedataGZ'+str(i)])
                    except:
                        continue
                #修改#----------赋予违法事实及自由裁量表格数据------
                for i in range(5):
                    try:
                        Illfacts_A= self.TextEdit_Illfact_listA[i]
                        Illfacts_A.setText(self.context['Face_illegfactA'+str(i)])
                        Illfacts_B= self.TextEdit_Illfact_listB[i]
                        Illfacts_B.setText(self.context['Face_illegfactB'+str(i)])
                        Illfacts_C= self.TextEdit_Illfact_listC[i]
                        Illfacts_C.setText(self.context['Face_illegfactC'+str(i)])
                        Illfacts_D= self.TextEdit_Illfact_listD[i]
                        Illfacts_D.setText(self.context['Face_illegfactD'+str(i)])                           
                        Illfacts_E= self.TextEdit_Illfact_listE[i]
                        Illfacts_E.setText(self.context['Face_illegfactE'+str(i)])                           
                    except:
                        continue
                #修改#----------赋予告知书违法事实及自由裁量表格数据------
                for i in range(3):
                    try:
                        Illfacts_GZA= self.TextEdit_Illfact_listGZA[i]
                        Illfacts_GZA.setText(self.context['Face_illegfactGZA'+str(i)])
                        Illfacts_GZB= self.TextEdit_Illfact_listGZB[i]
                        Illfacts_GZB.setText(self.context['Face_illegfactGZB'+str(i)])
                        Illfacts_GZC= self.TextEdit_Illfact_listGZC[i]
                        Illfacts_GZC.setText(self.context['Face_illegfactGZC'+str(i)])
                        Illfacts_GZD= self.TextEdit_Illfact_listGZD[i]
                        Illfacts_GZD.setText(self.context['Face_illegfactGZD'+str(i)])                           
                        Illfacts_GZE= self.TextEdit_Illfact_listGZE[i]
                        Illfacts_GZE.setText(self.context['Face_illegfactGZE'+str(i)])    
                    except:
                        continue
                #————————————————————————————————————————————————————————————————
                self.textEdit_Punishment_join.setText(self.context['Face_Punishment'])
                self.lineEdit_Presenter.setText(self.context['Case_Presenter'])
                self.lineEdit_Representor.setText(self.context['Face_Representor'])
                self.lineEdit_Otherperson.setText(self.context['Face_Otherperson'])
                self.lineEdit_Adress.setText(self.context['Case_Adress'])
                self.lineEdit_Heyitime.setText(self.context['Case_Heyitime'])
                self.lineEdit_refers.setText(self.context['Face_refers'])
                self.lineEdit_numdocGZ.setText(self.context['Case_NumGZ'])
                self.radioButton_Hearing.setChecked(self.context['Face_Hearing'])
                self.dateEdit_filingtimeR.setDate(self.context['Face_timeR'])
                self.lineEdit_linenumber.setText(self.context['Case_lineNUM'])
                self.lineEdit_linePerple.setText(self.context['Case_linePerple'])
                self.lineEdit_lineadrss.setText(self.context['Case_lineadrss'])
                self.lineEdit_linecode.setText(self.context['Case_linecode'])
                self.textEdit_absractcase_2.setText(self.context['Case_abstractGZ'])
                self.textEdit_PunishmentGZ.setText(self.context['Case_PunishmentGZ'])
                self.lineEdit_numdocJD.setText(self.context['Case_NumJD'])
                self.dateEdit_GZtime.setDate(self.context['Face_GZtime'])
                self.textEdit_PunishmentJD.setText(self.context['Face_PunishmentJD'])
                self.textEdit_informed.setText(self.context['Case_Informed'])
                self.lineEdit_execuWay.setText(self.context['Case_ExecuWay'])
                self.dateEdit_execuTime.setDate(self.context['Face_timeEX'])
                self.lineEdit_execuMode.setText(self.context['Case_ExecuMode'])
                self.textEdit_heyi.setText(self.context['Face_heyi'])
                self.textEdit_heyi2.setText(self.context['Face_heyi2'])
                self.textEdit_show.setText(self.context['Face_show'])
                self.dateEdit_JDTime.setDate(self.context['Face_timeJD'])
        #-----读取时候刷新段落格式——————————————
        textEdit_list=[self.textEdit_litigant,self.textEdit_absractcase,self.textEdit_absractcase_2,self.textEdit_process,self.textEdit_heyi,self.textEdit_heyi2,self.textEdit_informed,self.textEdit_show]
        self.blockFormat_initA(textEdit_list,True)
    #保存数据函数 #界面赋值给存储对象
    def SAVE_context(self,File_name):
        self.context['Case_name']=self.textEdit_litigant.toPlainText()
        self.context['Face_Accuname']=self.textEdit_accuss.toPlainText()
        self.AccuLists=self.context['Face_Accuname'].split('；')#会不会造成重复？
        Accuname = re.sub(r'[ ]*\d$', '', self.comboBox_accusDC.itemText(0))# 使用正则表达式去除最后一个数字以及它前面的所有连续空格（如果有）
        Accuname = Accuname.replace('案', '等案')  if len(self.AccuLists) > 1 else Accuname
        self.context['Case_accuname']=Accuname#输出案由为案由列表第一项文本
        self.context['Case_Source']=self.comboBox_source.currentText()
        self.context['Face_SourceIndex']=self.comboBox_source.currentIndex()
        self.context['Face_time']=self.dateEdit_filingtime.date()
        filingtime=self.dateEdit_filingtime.date()
        filingtime=filingtime.toString(Qt.DefaultLocaleLongDate)
        self.context['Case_time']=filingtime
        self.context['Case_abstract']=self.textEdit_absractcase.toPlainText()
        list_Law=[]
        list_Source=[]
        if self.AccuLists != []:
            for i in self.AccuLists:
                (listX, listY)=Accus_explode(i,self.Accus)
                list_Law += listX
                list_Source += listY
            (Caselaw, CaselawT)=Accus_join(list_Law)
            (Casesource, CasesourceT)=Accus_join(list_Source)
        else:
            (Caselaw, CaselawT, Casesource, CasesourceT)=('', '', '', '')
        self.context['Case_law']=Caselaw
        self.context['Case_lawsource']=Casesource
        self.context['Face_process']=self.textEdit_process.toPlainText()
        CaseProcess=self.context['Face_process'].replace('\n', '\a')
        self.context['Case_process']=Listing(CaseProcess)
        self.context['Face_heyi']=self.textEdit_heyi.toPlainText()
        heyi=self.context['Face_heyi'].replace('\n', '\a')
        self.context['Case_heyi']=Listing(heyi)
        self.context['Face_heyi2']=self.textEdit_heyi2.toPlainText()
        heyi2=self.context['Face_heyi2'].replace('\n', '\a')
        self.context['Case_heyi2']=Listing(heyi2)
        self.context['Face_AccuDCIndex']=self.comboBox_accusDC.currentIndex()
        #修改#---------违法事实自由裁量赋值-------------
        for i in range(5):
            try:
                Illfacts_A= self.TextEdit_Illfact_listA[i]
                self.context['Face_illegfactA'+str(i)]=Illfacts_A.toPlainText()
                Illfacts_B= self.TextEdit_Illfact_listB[i]
                self.context['Face_illegfactB'+str(i)]=Illfacts_B.toPlainText()
                Illfacts_C= self.TextEdit_Illfact_listC[i]
                self.context['Face_illegfactC'+str(i)]=Illfacts_C.toPlainText()
                Illfacts_D= self.TextEdit_Illfact_listD[i]
                self.context['Face_illegfactD'+str(i)]=Illfacts_D.toPlainText()                  
                Illfacts_E= self.TextEdit_Illfact_listE[i]
                self.context['Face_illegfactE'+str(i)]=Illfacts_E.toPlainText()
            except:
                continue        
        for (x, y, z) in zip(['（一）','（二）','（三）','（四）','（五）'], range(5), ['A','B','C','D','E' ]):
            try:
                accu=self.AccuLists[y]
                (list_Law, list_Source)=Accus_explode(accu,self.Accus)
                (law,lawtext)=Accus_join(list_Law)
                (source, sourcetext)=Accus_join(list_Source)
                self.context['Case_FactLaw'+str(y)]=x+self.context['Face_illegfact'+z+'0']+'该行为已经违反了%s的规定。'%law
                self.context['Case_FactSource'+str(y)]='对违法行为%s，应依据%s的规定,参照%s的裁量标准，考虑%s情节，建议予以当事人%s的行政处罚。'%(x, source,self.context['Face_illegfact'+z+'3'] ,self.context['Face_illegfact'+z+'1'],self.context['Face_illegfact'+z+'2'] )
                self.context['Case_discretionS'+str(y)]=self.context['Face_illegfact'+z+'3']+'：'+self.context['Face_illegfact'+z+'4']
            except:
                break
        #————————————————————————————————————————————————————————————————————#
        (Case_illeglFacts, Case_discretion, Case_discretionS)=('', '', '')

        if len(self.AccuLists)<=1:#单案由案件
            Case_illeglFacts=self.context['Face_illegfactA0']+'该行为已经违反了%s的规定。'%law
            punish=self.textEdit_Punishment_join.toPlainText()
            Case_discretion='对上述违法行为，应依据%s的规定,参照%s的裁量标准，考虑%s情节，建议予以当事人以下行政处罚：'%(source,self.context['Face_illegfactA'+'3'] ,self.context['Face_illegfactA'+'1'])+'\a'+punish.replace('\n', '\a')
            Case_discretionS=self.context['Face_illegfactA'+'3']+'：'+self.context['Face_illegfactA'+'4']
        else:
            for i  in range(len(self.AccuLists)):
                Case_illeglFacts += self.context['Case_FactLaw'+str(i)]+'\a'
                Case_discretion += self.context['Case_FactSource'+str(i)]+'\a'
                Case_discretionS += self.context['Case_discretionS'+str(i)]+'\a'
            Case_illeglFacts=Case_illeglFacts.strip('\a')
            punish=self.textEdit_Punishment_join.toPlainText()
            Case_discretion=Case_discretion.strip('\a')+'\a'+'经分别裁量，现合并上述处罚。建议予以当事人以下行政处罚：'+'\a'+punish.replace('\n', '\a')
            Case_discretionS=Case_discretionS.strip('\a')
        CaseLawS=CaselawT+'\a'+CasesourceT+'\a'+Case_discretionS
        self.Case_illeglFacts = Case_illeglFacts.replace('\u0007','\n')
        self.context['Case_LAWS']=Listing(CaseLawS)
        self.context['Case_illeglFacts']=Listing(Case_illeglFacts)
        self.context['Face_discretion']=Case_discretion.replace('\a', '')
        self.context['Case_discretion']=Listing(Case_discretion)
        self.context['Face_Punishment']=punish
        Evidencelist=[]
        Case_evidences=''
        #修改#---------证据组赋值-------------
        for (i, j) in zip(range(8), ['（一）','（二）','（三）','（四）','（五）','（六）','（七）','（八）']):
            try:
                Evidence = self.TextEdit_Evidence_list[i]
                self.context['Face_evidence'+str(i)]=Evidence.toPlainText()
                Evidenceproof = self.TextEdit_Evidenceproof_list[i]
                self.context['Face_evidencedata'+str(i)] = Evidenceproof.toPlainText()
                if self.context['Face_evidence'+str(i)]  != '' and  self.context['Face_evidencedata'+str(i)] !='':
                    Case_evidences += '证据组%s：'%j+self.context['Face_evidence'+str(i)] +self.context['Face_evidencedata'+str(i)]+'\a'
                    Evidence=self.context['Face_evidence'+str(i)]
                    Evidence=Evidence.replace('。', '')
                    Evidencelist.append(Evidence)
            except:
                continue
        #-------------------------------------------------#
        self.context['Case_evidenceAll']='；'.join(Evidencelist)+'。'
        self.context['Case_evidences']=Listing(Case_evidences.strip('\a')   )
        #————————合议记录
        self.context['Case_Presenter']=self.lineEdit_Presenter.text()
        self.context['Face_Representor']=self.lineEdit_Representor.text()
        self.context['Face_Otherperson']=self.lineEdit_Otherperson.text()
        self.context['Case_Persons']=f"{self.context['Case_Presenter']}、{self.context['Face_Otherperson']}、{self.context['Face_Representor']}"
        self.context['Case_Adress']=self.lineEdit_Adress.text()
        self.context['Case_Heyitime']=self.lineEdit_Heyitime.text()
        #————————事先告知
        self.context['Face_refers']=self.lineEdit_refers.text()
        self.context['Case_NumGZ']=self.lineEdit_numdocGZ.text()
        self.context['Face_Hearing']=self.radioButton_Hearing.isChecked()
        if self.context['Face_Hearing']:
            self.context['Case_Hearing']='拟对%s作出的行政处罚属于《中华人民共和国行政处罚法》第六十三条第一款规定范围，%s有权要求组织听证。根据《中华人民共和国行政处罚法》第六十四条第（一）项的规定，如%s要求组织听证，应当在收到本告知书后五日内提出申请，逾期视为放弃听证。'%(self.context['Face_refers'],self.context['Face_refers'], self.context['Face_refers'])
        else:
            self.context['Case_Hearing']=''
        self.context['Face_timeR']=self.dateEdit_filingtimeR.date()
        Rfilingtime=self.dateEdit_filingtimeR.date()
        Rfilingtime=Rfilingtime.toString(Qt.DefaultLocaleLongDate)
        self.context['Case_timeR']=Rfilingtime
        self.context['Case_lineNUM']=self.lineEdit_linenumber.text()
        self.context['Case_linePerple']=self.lineEdit_linePerple.text()
        self.context['Case_lineadrss']=self.lineEdit_lineadrss.text()
        self.context['Case_linecode']=self.lineEdit_linecode.text()
        self.context['Case_abstractGZ']=self.textEdit_absractcase_2.toPlainText()
        self.context['Face_AccuGZIndex']=self.comboBox_accusGZ.currentIndex()
        #修改#---------告知违法事实自由裁量赋值-------------
        for i in range(3):
            try:
                Illfacts_GZA= self.TextEdit_Illfact_listGZA[i]
                self.context['Face_illegfactGZA'+str(i)]=Illfacts_GZA.toPlainText()
                Illfacts_GZB= self.TextEdit_Illfact_listGZB[i]
                self.context['Face_illegfactGZB'+str(i)]=Illfacts_GZB.toPlainText()
                Illfacts_GZC= self.TextEdit_Illfact_listGZC[i]
                self.context['Face_illegfactGZC'+str(i)]=Illfacts_GZC.toPlainText()
                Illfacts_GZD= self.TextEdit_Illfact_listGZD[i]
                self.context['Face_illegfactGZD'+str(i)]=Illfacts_GZD.toPlainText()                  
                Illfacts_GZE= self.TextEdit_Illfact_listGZE[i]
                self.context['Face_illegfactGZE'+str(i)]=Illfacts_GZE.toPlainText()
            except:
                continue
        for (x, y, z) in zip(['（一）','（二）','（三）','（四）','（五）'], range(5), ['A','B','C','D','E' ]):
            try:
                accu=self.AccuLists[y]
                (list_Law, list_Source)=Accus_explode(accu,self.Accus)
                (law,lawtext)=Accus_join(list_Law)
                (source, sourcetext)=Accus_join(list_Source)
                self.context['Case_FactLawGZ'+str(y)]=x+self.context['Face_illegfactGZ'+z+'0']+'该行为已经违反了%s的规定。'%law
                self.context['Case_FactSourceGZ'+str(y)]='对违法行为%s，本机关依据%s的规定,参照%s的裁量标准，考虑%s情节，拟给予%s%s的行政处罚。'%(x, source,self.context['Face_illegfact'+z+'3'] ,self.context['Face_illegfactGZ'+z+'1'],self.context['Face_refers'], self.context['Face_illegfactGZ'+z+'2'] )
            except:
                break
        (Case_illeglFacts, Case_discretion)=('', '')
        if len(self.AccuLists)<=1:#单案由案件
            Case_illeglFacts=self.context['Face_illegfactGZA0']+'该行为已经违反了%s的规定。'%law
            punish=self.textEdit_PunishmentGZ.toPlainText()
            Case_discretion='对上述违法行为，本机关依据%s的规定,参照%s的裁量标准，考虑%s情节，拟给予%s以下行政处罚：'%(source,self.context['Face_illegfactA'+'3'] ,self.context['Face_illegfactGZA'+'1'],self.context['Face_refers'] )+'\a'+punish.replace('\n', '\a')
        else:
            for i  in range(len(self.AccuLists)):
                Case_illeglFacts += self.context['Case_FactLawGZ'+str(i)]+'\a'
                Case_discretion += self.context['Case_FactSourceGZ'+str(i)]+'\a'
            Case_illeglFacts=Case_illeglFacts.strip('\a')
            punish=self.textEdit_PunishmentGZ.toPlainText()
            Case_discretion=Case_discretion.strip('\a')+'\a'+'经分别裁量，现合并上述处罚。本机关拟给予%s以下行政处罚：'%self.context['Face_refers']+'\a'+punish.replace('\n', '\a')
        self.context['Case_illeglFactsGZ']=Listing(Case_illeglFacts)
        self.context['Case_discretionGZ']=Listing(Case_discretion)
        self.context['Case_PunishmentGZ']=punish
        Case_evidences=''
        for (i, j) in zip(range(8), ['（一）','（二）','（三）','（四）','（五）','（六）','（七）','（八）']):
            try:
                EvidenceGZ = self.TextEdit_EvidenceGZ_list[i]
                self.context['Face_evidenceGZ'+str(i)]=EvidenceGZ.toPlainText()
                EvidenceproofGZ = self.TextEdit_EvidenceproofGZ_list[i]
                self.context['Face_evidencedataGZ'+str(i)] = EvidenceproofGZ.toPlainText()
                if self.context['Face_evidenceGZ'+str(i)] !='' and self.context['Face_evidencedataGZ'+str(i)] !='' :
                    Case_evidences += '证据组%s：'%j+self.context['Face_evidenceGZ'+str(i)] +self.context['Face_evidencedataGZ'+str(i)]+'\a'
            except:
                break
        self.context['Case_evidencesGZ']=Listing(Case_evidences.strip('\a'))
        self.context['Case_NumJD']=self.lineEdit_numdocJD.text()
        self.context['Face_PunishmentJD']=self.textEdit_PunishmentJD.toPlainText()
        Case_discretionJD=Case_discretion.replace(punish,self.context['Face_PunishmentJD'].replace('\n', '\a'))
        Case_discretionJD=Case_discretionJD.replace('拟', '决定')
        self.context['Face_GZtime']=self.dateEdit_GZtime.date()
        self.context['Case_discretionJD']=Listing(Case_discretionJD)
        self.context['Case_Informed']=self.textEdit_informed.toPlainText()
        self.context['Face_timeJD']=self.dateEdit_JDTime.date()
        JDtime=self.dateEdit_JDTime.date()
        JDtime=JDtime.toString(Qt.DefaultLocaleLongDate)
        self.context['Face_show']=self.textEdit_show.toPlainText()
        showjd=self.context['Face_show'].replace('\n', '\a')
        self.context['Case_show']=Listing(showjd)
        self.context['Case_timeJD']=JDtime
        self.context['Case_ExecuWay']=self.lineEdit_execuWay.text()
        self.context['Face_timeEX']=self.dateEdit_execuTime.date()
        Executime=self.dateEdit_execuTime.date()
        Executime=Executime.toString(Qt.DefaultLocaleLongDate)
        self.context['Case_timeEX']=Executime
        self.context['Case_ExecuMode']=self.lineEdit_execuMode.text()
        with open (File_name,'wb') as self.context_file:
            dump(self.context, self.context_file)
            #self.signal_manager.signal.context_updated.emit(self.context)#文本更新信号发送
        print("保存成功")

    def extract_quantity(self,evidence_name):
        """
        从证据名称中提取数量
        """
        quantity_pattern = re.compile(r'(\d+)[份张页]|一份|一张|一页')
        match = quantity_pattern.search(evidence_name)
        if match:
            if match.group(1):  # 匹配到数字
                return int(match.group(1))
            else:  # 匹配到“一份”、“一张”等
                return 1
        return 1  # 默认数量为1
    
    def extract_date(self,evidence_name):
        """
        从证据名称中提取日期并转换为 YYYY-MM-DD 格式
        """
        date_pattern = re.compile(r'(\d{4})年(\d{1,2})月(\d{1,2})日')
        match = date_pattern.search(evidence_name)
        if match:
            year, month, day = match.groups()
            return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        return ""  # 如果没有日期，返回空字符串
    
    def clean_evidence_name(self,evidence_name):
        """
        清理证据名称：去除年月日和数量词，并将“照片”补充为“现场照片证据”
        """
        # 去除年月日
        #evidence_name = re.sub(r'\d{4}年\d{1,2}月\d{1,2}日', '', evidence_name)
        # 去除数量词（如“一份”、“一张”、“3份”、“2张”等）
        evidence_name = re.sub(r'\d+[份张页]|一份|一张|一页', '', evidence_name)
        # 将“照片”补充为“现场照片证据”
        evidence_name = evidence_name.replace("照片", "现场照片证据")
        # 去除多余的空格和标点
        evidence_name = re.sub(r'[；。\s]*$', '', evidence_name).strip()
        return evidence_name
    
    def clean_evidence_proof(self,evidence_proof):
        """
        清理证据证明：删除“本组证据证明”这六个字
        """
        return evidence_proof.replace("本组证据证明", "").strip()
    
    def split_compound_evidence(self,text):
        """
        拆分复合证据名称，支持以下两种情况：
        1. XXX、XXX各一份
        2. 日期、日期、日期询问笔录各一份
        """
        base_year=datetime.now().year
        #print(base_year+'基准年份')
    # 提取日期部分
        date_pattern = re.compile(r'((?:\d{4}年)?\d{1,2}月\d{1,2}日)')
        dates = []
        pos = 0
        
        while True:
            match = date_pattern.search(text, pos)
            if not match:
                break
            date_str = match.group(1)
            dates.append(date_str)
            pos = match.end()
            # 跳过分隔符
            if pos < len(text) and text[pos] in ['、', '，', ',', '以及', '和']:
                pos += 1
        
        # 处理基准年份
        base_year_used = next((int(re.search(r'(\d{4})年', d).group(1)) for d in dates if '年' in d), base_year)
        
        # 补全年份并格式化日期
        processed_dates = []
        for d in dates:
            if '年' not in d:
                processed_dates.append(f"{base_year_used}年{d}")
            else:
                processed_dates.append(d)
        
        # 提取证据描述部分
        evidence_desc = text[pos:].strip()
        
        # 分割证据类型
        split_pattern = re.compile(r'以及|及|和|、|，')
        evidence_parts = [p.strip() for p in split_pattern.split(evidence_desc) if p.strip()]
        
        result = []
        for part in evidence_parts:
            # 处理带有"各"的情况
            ge_match = re.match(r'^(.*?)\s*各\s*(\d+|[一二三四五六七八九十]+)\s*(\D+)$', part)
            if ge_match:
                evidence_type, amount, unit = ge_match.groups()
                for d in processed_dates:
                    result.append(f"{d}{evidence_type}{amount}{unit}")
                continue
            
            # 处理普通数量描述
            normal_match = re.match(r'^(.*?)\s*(\d+|[一二三四五六七八九十]+)\s*(\D+)$', part)
            if normal_match:
                evidence_type, amount, unit = normal_match.groups()
                for d in processed_dates:
                    result.append(f"{d}{evidence_type}{amount}{unit}")
        
        return result
    
    def type_evidence(self,evidence_name):
        evidence_type = "书证"
        if '现场笔录' in evidence_name:
            evidence_type = "勘验笔录、现场笔录"
        elif '检测报告' in  evidence_name:
            evidence_type = "鉴定意见"
        elif '询问笔录' in  evidence_name:
            evidence_type = "当事人的陈述"
        return evidence_type

    @pyqtSlot()
    def on_pushButton_Read_clicked(self):
        """
        Slot documentation goes here.
        """
        fname = QFileDialog.getOpenFileName(self, '打开文件','./Save/',(" 存档文件(*.pickle)"))
        if fname[0]:
            reply = QMessageBox.question(self,'询问','是否载入该存档？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
            if reply == QMessageBox.Yes:
                try:
                    self.READ_context('Default.pickle')
                except:
                    print('载入前清除失败，请检查Default.pickle文件。')
                try:
                    self.READ_context(fname[0])
                except:
                    QMessageBox.information(self,'错误提示', '载入失败！')
        # TODO: not implemented yet
        #raise NotImplementedError
#
    @pyqtSlot()
    def on_pushButton_Save_clicked(self):
        """
        Slot documentation goes here.
        """
        (SfileName, ok2 )= QFileDialog.getSaveFileName(self,  "文件保存",  './Save/',  "存档文件(*.pickle)")
        #print(SfileName, ok2)
        try:
            self.SAVE_context(SfileName)
            QMessageBox.information(self,'提示', '保存成功')
        except:
            QMessageBox.information(self,'提示', '未成功保存')
        # TODO: not implemented yet
        #raise NotImplementedError

    @pyqtSlot()
    def on_pushButton_Action_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        if self.textEdit_accuss.toPlainText() !='':
            self.SAVE_context('AutoSave.pickle')
        else:
            QMessageBox.information(self,'提示', '请先从案由列表内选择添加本案案由!')
            return
        litigant = self.textEdit_litigant.toPlainText()
        list_litigant=litigant.split('，')
        litigant_name=list_litigant[0]#获取当事人名称
        Doc_suffix=litigant_name
        self.MadeNote=''
        if self.checkBox_shouli.isChecked():
            print('制作受理立案文书……')
            #——————————文书替换————————
            THCH={'Case_Source','Case_name','Case_time','Case_abstract'}
            self.WORD_tpl('shoulijilu','受理记录',THCH, Doc_suffix)
            THCH={'Case_Source','Case_name','Case_time','Case_abstract', 'Case_law'}
            self.WORD_tpl('lianbaogao','立案报告', THCH, Doc_suffix)
        #调查终结报告
        if self.checkBox_zhjie.isChecked():
            print("制作调查终结报告文书……")
            THCH={'Case_name','Case_accuname','Case_process','Case_illeglFacts', \
            'Case_evidences','Case_discretion','Case_LAWS'}
            self.WORD_tpl('zhongjiebaogao','终结报告',THCH, Doc_suffix)
            #导出同步用的json文件
            
            caseFirst_Date = datetime.now().date().strftime("%Y-%m-%d")
            #print(caseFirst_Date)
            evidence_list = []
            for i in range(8):
                face_evidence = self.TextEdit_Evidence_list[i].toPlainText()
                face_evidencedata = self.TextEdit_Evidenceproof_list[i].toPlainText()
                if face_evidence and face_evidencedata:
                    evidence_items = re.split(r'[；，]', face_evidence)
                    group = []
                    for evidence in evidence_items:
                        if evidence.strip():  # 去掉空行
                            # 拆分复合证据名称
                            #print(evidence)
                            split_evidence = self.split_compound_evidence(evidence)
                            for sub_evidence in split_evidence:
                                if sub_evidence.strip():  # 去掉空行
                                    # 提取数量和日期
                                    quantity = self.extract_quantity(sub_evidence)
                                    collection_date = self.extract_date(sub_evidence)
                                    # 清理证据名称
                                    evidence_name = self.clean_evidence_name(sub_evidence)
                                    # 构建证据条目
                                    # 根据证据条目来测定证据类型
                                    evidence_type = self.type_evidence(evidence_name)# 默认类型为书证
                                    if '笔录' in evidence_name or '报告' in evidence_name:
                                        caseFirst_Date = compare_dates(collection_date,caseFirst_Date)                                       
                                    evidence_entry = {
                                        "evidenceName": evidence_name,
                                        "quantity": quantity,
                                        "location": "",
                                        "description": "",
                                        "collector": "",
                                        "evidenceType": evidence_type,  # 默认类型为书证
                                        "collectionDate": collection_date
                                    }
                                    # 只在第一条证据中添加 evidenceProof
                                    if not group:  # 如果是组内的第一条证据
                                        evidence_entry["evidenceProof"] = self.clean_evidence_proof(face_evidencedata.strip())
                                        # if not evidence_list: # 如果是第一组证据
                                        #     evidence_entry["caseSummary"] =self.textEdit_absractcase.toPlainText()
                                        #     evidence_entry["caseProcess"] =self.textEdit_process.toPlainText()
                                        #     evidence_entry["caseFacts"] = self.Case_illeglFacts
                                        #     evidence_entry["caseFirstDate"] = "2025-02-05"
                                    group.append(evidence_entry)

                    # 如果组不为空，则添加到证据列表中
                    if group:
                        evidence_list.append(group)
            #在组内的第一个列表的第一个字典中追加添加数据
            if evidence_list[0]:
                evidence_first = evidence_list[0][0]
                #print(evidence_first)
                evidence_first["caseSummary"] =self.textEdit_absractcase.toPlainText()
                evidence_first["caseProcess"] =self.textEdit_process.toPlainText()
                evidence_first["caseFacts"] = self.Case_illeglFacts
                evidence_first["caseFirstDate"] = caseFirst_Date
                evidence_list[0][0] = evidence_list[0][0]
                
            # 保存为 JSON 文件
            output_file = r"Output/%s/output.json"%Doc_suffix
            with open(output_file, 'w', encoding='utf-8') as json_file:
                json.dump(evidence_list, json_file, ensure_ascii=False, indent=4)

            print(f"调查报告信息提取完成并保存为 {output_file} 文件！")
                
            #增加导出,现值移入历史值
            FT_History={key:value for key, value in self.History.items()if key == 'Fst_History'}
            DicFst_History=dict(FT_History['Fst_History'])#空值出错
            if DicFst_History == {}:
                DicFst_History['当事人'] = ''
            if self.textEdit_litigant.toPlainText() != '' and self.textEdit_litigant.toPlainText() != DicFst_History['当事人'] :
                self.History['Tre_History'] = self.History['Sed_History']
                self.History['Sed_History'] = self.History['Fst_History']
                #DicFst_History={}
                if self.textEdit_litigant.toPlainText() != '':DicFst_History['当事人'] = self.textEdit_litigant.toPlainText()
                if self.textEdit_absractcase.toPlainText() != '':DicFst_History['案情摘要'] = self.textEdit_absractcase.toPlainText()
                if self.textEdit_process.toPlainText() != '':DicFst_History['调查过程'] = self.textEdit_process.toPlainText()
                #修改# 现值写入历史值——————————————————————————————————————
                for (i, j) in zip(range(8), ['一','二','三','四','五','六','七','八']):
                    Evidence = self.TextEdit_Evidence_list[i]
                    DicFst_History['证明材料%s'%j] = Evidence.toPlainText() if Evidence.toPlainText() !='' else ''
                    Evidenceproof = self.TextEdit_Evidenceproof_list[i]
                    DicFst_History['证明内容%s'%j] = Evidenceproof.toPlainText() if Evidenceproof.toPlainText() !='' else ''
                tablelist =[self.TextEdit_Illfact_listA,self.TextEdit_Illfact_listB,self.TextEdit_Illfact_listC,self.TextEdit_Illfact_listD,self.TextEdit_Illfact_listE]
                for i in range(5):
                    print(i,self.comboBox_accusDC.count())
                    if i <= self.comboBox_accusDC.count():
                        X = self.comboBox_accusDC.itemText(i)
                        Illfactslist = tablelist[i]
                        #print(Illfactslist[0].toPlainText())
                        DicFst_History['%s违法事实'%X]=Illfactslist[0].toPlainText()if Illfactslist[0].toPlainText() != '' else ''
                        DicFst_History['%s裁量情节'%X]=Illfactslist[1].toPlainText()if Illfactslist[1].toPlainText() != '' else ''
                        DicFst_History['%s处罚建议'%X]=Illfactslist[2].toPlainText()if Illfactslist[2].toPlainText() != '' else ''
                        DicFst_History['%s标准内容'%X]=Illfactslist[4].toPlainText()if Illfactslist[4].toPlainText() != '' else ''
                    else:
                        continue
                self.History['Fst_History'] = DicFst_History
            with open ('History.pickle', 'wb') as History_file :
                dump(self.History, History_file)
        if self.checkBox_heyi.isChecked():
            print("制作合议记录……")
            self.context['Case_accuname2']=litigant_name+self.context['Case_accuname']
            THCH={'Case_accuname2','Case_Presenter','Case_Persons','Case_Heyitime',\
            'Case_Adress','Case_heyi','Case_heyi2'}
            self.WORD_tpl('heyijilu','合议记录',THCH, Doc_suffix)
        if self.checkBox_gaozhi.isChecked():
            print("制作事先告知书文书……")
            THCH={'Case_NumGZ','Case_name','Case_abstractGZ','Case_timeR','Case_illeglFactsGZ',\
            'Case_evidencesGZ','Case_discretionGZ', 'Case_LAWS', 'Face_refers', 'Case_Hearing', 'Case_linecode', 'Case_lineNUM', \
            'Case_lineadrss', 'Case_linePerple'}
            self.WORD_tpl('shixiangaozhi','事先告知',THCH, Doc_suffix)
            print("制作事先告知书审批表文书……")
            THCH={'Case_name', 'Case_accuname', 'Case_illeglFacts','Case_evidenceAll' , 'Case_law','Case_lawsource' , 'Case_PunishmentGZ'}
            self.WORD_tpl('shenpi_gaozhi','事先告知审批表',THCH, Doc_suffix)
        if self.checkBox_chufa.isChecked():
            print("制作处罚决定书文书……")
            THCH={'Case_NumJD','Case_name','Case_abstractGZ','Case_timeR','Case_illeglFactsGZ','Case_law', \
            'Case_evidencesGZ','Case_discretionJD','Case_LAWS', 'Face_refers', 'Case_Informed'}
            self.WORD_tpl('chufajueding','处罚决定',THCH, Doc_suffix)
            print("制作决定书审批表文书……")
            THCH={'Case_name', 'Case_accuname', 'Case_illeglFacts','Case_evidenceAll' , 'Case_law','Case_lawsource' , 'Face_PunishmentJD'}
            self.WORD_tpl('shenpi_jueding','决定书审批表',THCH, Doc_suffix)
        if self.checkBox_gongshi.isChecked():
            print("制作决定书公示摘要……")
            self.context['Case_name2']=litigant_name
            THCH={'Case_NumJD','Case_name2','Case_timeJD', 'Case_show'}
            self.WORD_tpl('juedinggongshi','决定书公示摘要',THCH, Doc_suffix)
        if self.checkBox_jiean.isChecked():
            print("制作结案报告……")
            THCH={'Case_name', 'Case_Source', 'Case_timeR','Case_accuname', 'Case_NumJD' , 'Case_ExecuWay', 'Case_timeEX', 'Case_ExecuMode','Face_PunishmentJD'}
            self.WORD_tpl('jieanbaogao','结案报告',THCH, Doc_suffix)
        if self.MadeNote != '':
            QMessageBox.information(self,'提示', self.MadeNote+'请到Output文件夹查看。')
            reply = QMessageBox.question(self,'询问','是否打开Output文件夹', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
            if reply == QMessageBox.Yes:
                path_output = getcwd()
                path_output=path.join(path_output, 'Output')
                open_folder(path_output)
        else:
            QMessageBox.information(self,'提示', '请勾选需要制作的文书。')
    @pyqtSlot(QListWidgetItem)
    def on_listWidget_AccuEdit_itemClicked(self, item):
        """
        Slot documentation goes here.
        @param item DESCRIPTION
        @type QListWidgetItem
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        Accuname = self.lineEdit_Accus.text()
        if  Accuname:
            (list_Law, list_Source)=Accus_explode(Accuname,self.Accus)
            for(i,k, j)  in zip(list_Law[0::2], list_Law[1::2], range(5)):
                datai=QTableWidgetItem(i) #转换后可插入表格
                self.tableWidget_laws.setItem(j,0,datai)
                datak=QTableWidgetItem(k)
                self.tableWidget_laws.setItem(j,1,datak)
            for(i,k, j)  in zip(list_Source[0::2], list_Source[1::2], range(5)):
                datai=QTableWidgetItem(i) #转换后可插入表格
                self.tableWidget_lawsources.setItem(j,0,datai)
                datak=QTableWidgetItem(k)
                self.tableWidget_lawsources.setItem(j,1,datak)

    @pyqtSlot(QListWidgetItem)
    def on_listWidget_AccuEdit_itemDoubleClicked(self, item):

        """
        双击删除案由，并且清空相应的表格内的内容
        Slot documentation goes here.
        @param item DESCRIPTION
        @type QListWidgetItem
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        if self.tabWidget.currentIndex() == 0 and self.tabWidget_5.currentIndex() == 0:
            text_list = self.listWidget_AccuEdit.selectedItems()
            text =text_list[0].text()
            if text in self.AccuLists:
                reply = QMessageBox.question(self,'询问','该案由已经添加，是否确定删除？删除该案由会同时删除该案由下对应内容', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
                if reply == QMessageBox.Yes:
                    i=self.comboBox_accusDC.findText(text)
                    #修改#删除案由时，轮流对违法事实自由裁量重新赋值
                    for x in range(5):
                        Illfacts_A= self.TextEdit_Illfact_listA[x]
                        Illfacts_B= self.TextEdit_Illfact_listB[x]
                        Illfacts_C= self.TextEdit_Illfact_listC[x]
                        Illfacts_D= self.TextEdit_Illfact_listD[x]
                        Illfacts_E= self.TextEdit_Illfact_listE[x]
                        if i ==0:
                            Illfacts_A.setText(Illfacts_B.toPlainText())
                            Illfacts_B.setText(Illfacts_C.toPlainText())
                            Illfacts_C.setText(Illfacts_D.toPlainText())
                            Illfacts_D.setText(Illfacts_E.toPlainText())                       
                        elif i ==1:
                            Illfacts_B.setText(Illfacts_C.toPlainText())
                            Illfacts_C.setText(Illfacts_D.toPlainText())
                            Illfacts_D.setText(Illfacts_E.toPlainText())
                        elif i ==2:
                            Illfacts_C.setText(Illfacts_D.toPlainText())
                            Illfacts_D.setText(Illfacts_E.toPlainText())                           
                        elif i ==3:
                            Illfacts_D.setText(Illfacts_E.toPlainText())
                        Illfacts_E.setText('')
                    #————————————————————————————————————————————————————————————————————————————————————————————————
                    self.AccuLists.remove(text)
                    self.on_pushButton_Joinpunish_clicked()#重新合并处罚建议
                    QMessageBox.information(self,'提示', '终结报告中该案由相关内容已删除。\n请及时修改合议、事先告知等文书的相关内容！')
            elif len(self.AccuLists)<5 :#当案由数量少于5 ，并且新案由双击为添加
                self.AccuLists.append(text)
            else:
                QMessageBox.information(self,'提示', '软件最多只支持五个案由！')
            while '' in self.AccuLists:
                self.AccuLists.remove('')
            accuss='；'.join(self.AccuLists)
            accuss=accuss.strip('；')
            self.textEdit_accuss.setText(accuss)
            self.comboBox_accusGZ.clear()
            self.comboBox_accusDC.clear()
            self.comboBox_accusDC.addItems(self.AccuLists)
            self.comboBox_accusGZ.addItems(self.AccuLists)
        else:
             QMessageBox.information(self,'提示', '添加删除该案案由只能在文书制作页面1（立案界面）完成！')
    @pyqtSlot()
    def on_pushButton_accuAdd_clicked(self):
        """
        Slot documentation goes here.
        新增案由按钮
        """
        reply = QMessageBox.question(self,'询问','是否确定新增或修改该案由？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
        if reply == QMessageBox.Yes:
            with open ('./Date/Accus0.pickle', 'wb') as Accus_file :
                dump(self.Accus, Accus_file)
            Accuname = self.lineEdit_Accus.text()
            AccuLaw =''
            AccuSource =''
            if Accuname !="":
                for i in range(5):
                    for j in range(2):
                        if self.tableWidget_laws.item(i, j) != None and self.tableWidget_laws.item(i, j).text() !='':
                            datalaw = self.tableWidget_laws.item(i, j).text()
                        else :
                            break
                        AccuLaw += datalaw+'|'
                for i in range(5):
                    for j in range(2):
                        if self.tableWidget_lawsources.item(i, j) != None and self.tableWidget_lawsources.item(i, j).text() !='':
                            datasource=self.tableWidget_lawsources.item(i, j).text()
                        else:
                            break
                        AccuSource += datasource+'|'
                self.Accus[Accuname]={'违反条款': AccuLaw, '依据条款': AccuSource}
                try:
                    (listX, listY)=Accus_explode(Accuname,self.Accus)
                    #print(listX[0::2], listY[0::2])
                    Accus_join(listX)
                    Accus_join(listY)
                    accuAdd = True
                except:
                    QMessageBox.information(self,'错误提示', '案由格式错误，请参照已添加案由填写，并注意《》！')
                    accuAdd = False
                if accuAdd :
                    with open ('Accus.pickle', 'wb') as Accus_file :
                        dump(self.Accus, Accus_file)
                    self.Accus_initA()
            else :
                QMessageBox.information(self,'错误提示', '请输入案由名称以及相关内容！')

        # TODO: not implemented yet
        #raise NotImplementedError
    @pyqtSlot()
    def on_pushButton_accuCtrlV_clicked(self):
        """
        Slot documentation goes here.
        """
        (SfileName, ok2 )= QFileDialog.getSaveFileName(self,  "文件保存",  './Date/',  "存档文件(*.pickle)")
        #print(SfileName, ok2)
        try:
            with open (SfileName, 'wb') as Accus_file :
                dump(self.Accus, Accus_file)
            QMessageBox.information(self,'提示', '导出成功')
        except:
            QMessageBox.information(self,'提示', '导出失败')
        # TODO: not implemented yet
        #raise NotImplementedError
    @pyqtSlot()
    def on_pushButton_accuDel_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        Accuname = self.lineEdit_Accus.text()
        if Accuname in self.AccuLists:
            QMessageBox.information(self,'错误提示', '无法删除正在使用的案由，请先在文书制作中删除！')
        else:
            reply = QMessageBox.question(self,'询问','是否确定删除该案由？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
            if reply == QMessageBox.Yes:
                if Accuname in self.Accus:
                    with open ('./Date/Accus0.pickle', 'wb') as Accus_file :
                        dump(self.Accus, Accus_file)
                    del self.Accus[Accuname]
                    with open ('Accus.pickle', 'wb') as Accus_file :
                        dump(self.Accus, Accus_file)
                    self.Accus_initA()
                else:
                    QMessageBox.information(self,'错误提示', '请从案由列表选择一个案由删除')
    @pyqtSlot()
    def on_pushButton_accuEXO_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        #导入前先备份
        fname = QFileDialog.getOpenFileName(self, '导入文件','./Date',(" 案由文件(*.pickle)"))
        if fname[0]:
            reply = QMessageBox.question(self,'询问','是否导入新案由文件%s？'%fname[0], QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
            if reply == QMessageBox.Yes:
                with open ('./Date/Accus0.pickle', 'wb') as Accus_file :
                    dump(self.Accus, Accus_file)
                with open (fname[0], 'rb') as Accus_file :
                    Accus2 = load(Accus_file)
                if 'Case_name' in Accus2 or 'Fst_History' in Accus2 :
                    QMessageBox.information(self,'错误提示', '该文件为案件非案由存档,无法导入！')
                else:
                    reply = QMessageBox.question(self,'询问','是否替换目前已有案由？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
                    if reply == QMessageBox.Yes:#覆盖
                        Accusnew={**self.Accus, **Accus2}
                    else:#不覆盖
                        Accusnew={**Accus2, **self.Accus}
                    with open ('Accus.pickle', 'wb') as Accus_file :
                        dump(Accusnew, Accus_file)
                    self.Accus_initA()
    @pyqtSlot()
    def on_pushButton_accuRdc_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        fname = QFileDialog.getOpenFileName(self, '还原文件','./Date',(" 案由文件(*.pickle)"))
        if fname[0]:
            reply = QMessageBox.question(self,'询问','是否还原该备份文件%s？'%fname[0], QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
            if reply == QMessageBox.Yes:
                try:
                    with open (fname[0], 'rb') as Accus_file :
                       Accus0 = load(Accus_file)
                    if 'Case_name' in Accus0 or 'Fst_History' in Accus0 :
                        QMessageBox.information(self,'错误提示', '该文件非案由存档,无法还原！')
                    else:
                        with open ('Accus.pickle', 'wb') as Accus_file :
                           dump(Accus0, Accus_file)
                        self.Accus_initA()
                except:
                    QMessageBox.information(self,'错误提示', '还原失败，请查看Accus0备份文件是否被删除或破坏')
    #pushButton_accuTop
    @pyqtSlot()
    def on_pushButton_accuTop_clicked(self):
        """
        Slot documentation go here.

        """
        Accuname = self.lineEdit_Accus.text()
        if Accuname in self.Accus:
            Accus2={Accuname:self.Accus[Accuname]}
            Accusnew={**Accus2, **self.Accus}
            with open ('Accus.pickle', 'wb') as Accus_file :
                    dump(Accusnew, Accus_file)
            self.Accus_initA()
        else:
            QMessageBox.information(self,'错误提示', '请从案由列表选择一个案由。')

    @pyqtSlot()
    def on_pushButton_Clear_clicked(self):
        """
        Slot documentation goes here.

        """
        reply = QMessageBox.question(self,'询问','是否清除所有内容', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
        if reply == QMessageBox.Yes:

            try:
                self.READ_context('Default.pickle')
            except:
                QMessageBox.information(self,'错误提示', '清除失败，请查看Default文件是否被删除或破坏')
            try:
                self.dateEdit_execuTime.setDate(QDate.currentDate())
                self.dateEdit_GZtime.setDate(QDate.currentDate())
                self.dateEdit_filingtimeR.setDate(QDate.currentDate())
                self.dateEdit_filingtime.setDate(QDate.currentDate())
                self.dateEdit_JDTime.setDate(QDate.currentDate())
                self.SAVE_context('Default.pickle')
                DS1_History={key:value for key, value in self.History.items()if key == 'DeSetone'}
                DS2_History={key:value for key, value in self.History.items()if key == 'DeSettwo'}
                DicDS1_History=dict(DS1_History['DeSetone'])
                DicDS2_History=dict(DS2_History['DeSettwo'])
                self.lineEdit_Presenter.setText(DicDS1_History['Case_Presenter'])
                self.lineEdit_Representor.setText(DicDS1_History['Face_Representor'])
                self.lineEdit_Otherperson.setText(DicDS1_History['Face_Otherperson'])
                self.lineEdit_Adress.setText(DicDS1_History['Case_Adress'])
                self.lineEdit_linenumber.setText(DicDS2_History['Case_lineNUM'])
                self.lineEdit_linePerple.setText(DicDS2_History['Case_linePerple'])
                self.lineEdit_lineadrss.setText(DicDS2_History['Case_lineadrss'])
                self.lineEdit_linecode.setText(DicDS2_History['Case_linecode'])
            except:
                pass
        # TODO: not implemented yet
        #raise NotImplementedError
    @pyqtSlot()
    def on_pushButton_import_clicked(self):
        """
        Slot documentation goes here.
        """
        reply = QMessageBox.question(self,'询问','是否已经输入当事人统一指代称谓和立案时间', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
        if reply == QMessageBox.Yes and self.lineEdit_refers.text() !='':
            #print('前文内容导入中')
            self.SAVE_self.context('AutoSave.pickle')
            refers=self.context['Face_refers']
            self.context['Case_abstractGZ']=self.context['Case_abstract'].replace('当事人',refers)
            #修改#调查报告值传入事先告知——————————————————————————————————
            for i in range(8):
                try:
                    self.context['Face_evidenceGZ'+str(i)]=self.context['Face_evidence'+str(i)].replace('当事人',refers)
                    self.context['Face_evidencedataGZ'+str(i)]=self.context['Face_evidencedata'+str(i)].replace('当事人',refers)
                    EvidenceGZ = self.TextEdit_EvidenceGZ_list[i]
                    EvidenceGZ.setText(self.context['Face_evidenceGZ'+str(i)])
                    EvidenceproofGZ = self.TextEdit_EvidenceproofGZ_list[i]
                    EvidenceproofGZ.setText(self.context['Face_evidencedataGZ'+str(i)])
                except:
                    break
            for i in range(3):
                try:
                    self.context['Face_illegfactGZA'+str(i)]=self.context['Face_illegfactA'+str(i)].replace('当事人',refers)
                    Illfacts_GZA= self.TextEdit_Illfact_listGZA[i]
                    Illfacts_GZA.setText(self.context['Face_illegfactGZA'+str(i)])
                    
                    self.context['Face_illegfactGZB'+str(i)]=self.context['Face_illegfactB'+str(i)].replace('当事人',refers)
                    Illfacts_GZB= self.TextEdit_Illfact_listGZB[i]
                    Illfacts_GZB.setText(self.context['Face_illegfactGZB'+str(i)])
                    
                    self.context['Face_illegfactGZC'+str(i)]=self.context['Face_illegfactC'+str(i)].replace('当事人',refers)
                    Illfacts_GZC= self.TextEdit_Illfact_listGZC[i]
                    Illfacts_GZC.setText(self.context['Face_illegfactGZC'+str(i)])
                    
                    self.context['Face_illegfactGZD'+str(i)]=self.context['Face_illegfactD'+str(i)].replace('当事人',refers)
                    Illfacts_GZD= self.TextEdit_Illfact_listGZD[i]
                    Illfacts_GZD.setText(self.context['Face_illegfactGZD'+str(i)])
                    
                    self.context['Face_illegfactGZE'+str(i)]=self.context['Face_illegfactE'+str(i)].replace('当事人',refers)       
                    Illfacts_GZE= self.TextEdit_Illfact_listGZE[i]
                    Illfacts_GZE.setText(self.context['Face_illegfactGZE'+str(i)])
                except:
                    continue
            #——————————————————————————————————————————————————————————————————————————————————————————————————————
            self.context['Case_PunishmentGZ']=self.context['Face_Punishment'].replace('当事人',refers)
            self.textEdit_absractcase_2.setText(self.context['Case_abstractGZ'])
            self.textEdit_PunishmentGZ.setText(self.context['Case_PunishmentGZ'])
        elif self.lineEdit_refers.text() =="":
            QMessageBox.information(self,'错误提示', '请先输入当事人统一指代称谓和立案时间。')
        #raise NotImplementedError
    @pyqtSlot()
    def on_pushButton_import_2_clicked(self):
        """
        Slot documentation goes here.
        """
        self.SAVE_context('AutoSave.pickle')
        refers=self.context['Face_refers']
        GZtime=self.dateEdit_GZtime.date()
        GZtime=GZtime.toString(Qt.DefaultLocaleLongDate)
        self.context['Case_NumJD']=self.context['Case_NumGZ'].replace('告','')
        self.lineEdit_numdocJD.setText(self.context['Case_NumJD'])
        self.context['Face_PunishmentJD']=self.context['Case_PunishmentGZ']
        self.textEdit_PunishmentJD.setText(self.context['Face_PunishmentJD'])
        if self.context['Face_Hearing'] :
            self.textEdit_informed.setText('%s，本机关向%s直接送达了行政处罚事先告知书（%s），依法告知%s本机关拟作出行政处罚的事实、理由和依据，以及具体处罚内容和%s依法享有的陈述、申辩和听证的权利。%s在法定期限内未提出陈述、申辩以及听证。'%(GZtime, refers,self.context['Case_NumGZ'], refers, refers, refers))
        else:
            self.textEdit_informed.setText('%s，本机关向%s直接送达了行政处罚事先告知书（%s），依法告知%s本机关拟作出行政处罚的事实、理由和依据，以及具体处罚内容和%s依法享有的陈述、申辩的权利。%s在法定期限内未提出陈述、申辩。'%(GZtime,refers,self.context['Case_NumGZ'], refers, refers, refers))
    @pyqtSlot()
    def on_pushButton_import_heyi_clicked(self):
        """
        Slot documentation goes here.
        """
        litigant = self.textEdit_litigant.toPlainText()
        list_litigant=litigant.split('，')
        litigant_name=list_litigant[0]#获取当事人名称
        self.SAVE_context('AutoSave.pickle')
        HeyiPunishment =self.context['Face_Punishment'].replace('\n', '')
        self.context['Case_accuname2']=litigant_name+self.context['Case_accuname']
        heyitext=(f"{self.context['Case_Presenter']}：根据《卫生行政处罚程序》第二十五条规定，今天由{self.context['Case_Persons']}组成合议小组，对{self.context['Case_accuname2']}进行合议。请问是否有人需申请回避？如有，请现在提出。\n{self.context['Face_Representor']}、{self.context['Face_Otherperson']}：我与当事人没有利害冲突，不存在回避情形。\n{self.context['Case_Presenter']}：我也不存在回避情形。现没有人申请回避，合议小组正式组成。下面我们开始合议，首先由本案承办人员{self.context['Face_Representor']}介绍案情，并提出处理意见。\n{self.context['Face_Representor']}：{self.context['Face_process']}\n现已查明，当事人存在以下违法行为：\n{self.context['Case_illeglFacts']}\n主要证据有：{self.context['Case_evidenceAll']}\n我们承办人员认为：{self.context['Face_discretion']}\n{self.context['Case_Presenter']}：该案的基本情况已经了解，现在请合议人员就主体认定、违法事实、证据材料、法律适用、自由裁量、处罚程序等方面进行合议。\n\n{self.context['Case_Presenter']}、{self.context['Face_Representor']}、{self.context['Face_Otherperson']}\n：我认为本案违法主体认定清楚明确，违法事实由询问笔录、现场笔录等证据充分证明。\n：在处罚程序上面，我认真审查了该案件的材料，我们执法人员做到了取证合法有效，告知相对人权利义务正确，因此处罚程序合法正当。\n：在法律适用上，适用条款正确无误的。\n：自由裁量上，我们已经充分考虑当事人的相关情节，符合过罚相当的原则。\n\n{self.context['Case_Presenter']}：我认真审阅了案卷，同意大家的意见，对本案的主体、违法事实、裁量和法律适用等并无异议。大家还有什么需要补充吗？\n{self.context['Face_Representor']}、{self.context['Face_Otherperson']}：没有了。\n{self.context['Case_Presenter']}：现在就本次合议总结如下：本案的违法主体认定准确，违法事实清楚明确，证据材料确凿详实，处罚程序合法正当，自由裁量过罚相当，法律条款适用正确，拟给予当事人以下行政处罚:{HeyiPunishment}请大家表态。\n{self.context['Face_Representor']}、{self.context['Face_Otherperson']}：无异议。\n{self.context['Case_Presenter']}：我也同意。本案合议到此结束，请大家在合议记录上签字确认后离开。")
        heyitext=heyitext.replace('</w:t></w:r></w:p><w:p><w:r><w:t xml:space="preserve">', '\n')
        self.textEdit_heyi.setText(heyitext)
        heyiResult=(f"{litigant_name}存在以下违法行为：\n{self.context['Case_illeglFacts']}\n{self.context['Face_discretion']}")
        heyiResult=heyiResult.replace('</w:t></w:r></w:p><w:p><w:r><w:t xml:space="preserve">', '\n')
        self.textEdit_heyi2.setText(heyiResult)
    @pyqtSlot()
    def on_pushButton_import_show_clicked(self):
        """
        Slot documentation goes here.
        """
        self.SAVE_context('AutoSave.pickle')
        Case_evidenceAll=self.context['Case_evidenceAll'].replace('当事人', self.context['Face_refers'])
        showtext=(f"{self.context['Face_refers']}经查明存在以下违法事实：\n{self.context['Case_illeglFactsGZ']}\n主要证据有：{Case_evidenceAll}\n{self.context['Case_Informed']}\n{self.context['Case_discretionJD']}")
        showtext=showtext.replace('</w:t></w:r></w:p><w:p><w:r><w:t xml:space="preserve">', '\n')
        self.textEdit_show.setText(showtext)
    @pyqtSlot()
    def on_pushButton_Joinpunish_clicked(self):
        """
        Slot documentation goes here.
        """
        #修改 合并处罚内容

        text = self.textEdit_illfacts_A2.toPlainText()+'\n'+self.textEdit_illfacts_B2.toPlainText()+'\n'+\
            self.textEdit_illfacts_C2.toPlainText()+'\n'+self.textEdit_illfacts_D2.toPlainText()+'\n'+\
                self.textEdit_illfacts_E2.toPlainText()
        text=text.strip('\n')
        self.textEdit_Punishment_join.setText(text)
    #————————————————————历史输入模块————————————————————
    @pyqtSlot()
    def on_textEdit_litigant_selectionChanged(self):
        self.History_initA('当事人')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_absractcase_selectionChanged(self):
        self.History_initA('案情摘要')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_process_selectionChanged(self):
        self.History_initA('调查过程')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidencelist_0_selectionChanged(self):
        self.History_initA('证明材料一')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidencelist_1_selectionChanged(self):
        self.History_initA('证明材料二')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidencelist_2_selectionChanged(self):
        self.History_initA('证明材料三')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidencelist_3_selectionChanged(self):
        self.History_initA('证明材料四')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidencelist_4_selectionChanged(self):
        self.History_initA('证明材料五')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidencelist_5_selectionChanged(self):
        self.History_initA('证明材料六')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidencelist_6_selectionChanged(self):
        self.History_initA('证明材料七')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidencelist_7_selectionChanged(self):
        self.History_initA('证明材料八')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidenceproof_0_selectionChanged(self):
        self.History_initA('证明内容一')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidenceproof_1_selectionChanged(self):
        self.History_initA('证明内容二')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidenceproof_2_selectionChanged(self):
        self.History_initA('证明内容三')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidenceproof_3_selectionChanged(self):
        self.History_initA('证明内容四')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidenceproof_4_selectionChanged(self):
        self.History_initA('证明内容五')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidenceproof_5_selectionChanged(self):
        self.History_initA('证明内容六')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidenceproof_6_selectionChanged(self):
        self.History_initA('证明内容七')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_evidenceproof_7_selectionChanged(self):
        self.History_initA('证明内容八')
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_A0_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(0)
        self.History_initA('%s违法事实'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_A1_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(0)
        self.History_initA('%s裁量情节'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_A2_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(0)
        self.History_initA('%s处罚建议'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_A4_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(0)
        self.History_initA('%s标准内容'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_B0_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(1)
        self.History_initA('%s违法事实'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_B1_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(1)
        self.History_initA('%s裁量情节'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_B2_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(1)
        self.History_initA('%s处罚建议'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_B4_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(1)
        self.History_initA('%s标准内容'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_C0_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(2)
        self.History_initA('%s违法事实'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_C1_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(2)
        self.History_initA('%s裁量情节'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_C2_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(2)
        self.History_initA('%s处罚建议'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_C4_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(2)
        self.History_initA('%s标准内容'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_D0_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(3)
        self.History_initA('%s违法事实'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_D1_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(3)
        self.History_initA('%s裁量情节'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_D2_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(3)
        self.History_initA('%s处罚建议'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_D4_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(3)
        self.History_initA('%s标准内容'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_E0_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(4)
        self.History_initA('%s违法事实'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_E1_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(4)
        self.History_initA('%s裁量情节'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_E2_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(4)
        self.History_initA('%s处罚建议'%X)
        self.tabWidge_list.setCurrentIndex(0)
    @pyqtSlot()
    def on_textEdit_illfacts_E4_selectionChanged(self):
        X = self.comboBox_accusDC.itemText(4)
        self.History_initA('%s标准内容'%X)
        self.tabWidge_list.setCurrentIndex(0)
    
    #————————————————————————————————————————————————————————————————————————————————————————————
    
    @pyqtSlot()
    def on_pushButton_Fordefault1_clicked(self):
        """
        Slot documentation goes here.
        """
        reply = QMessageBox.question(self,'询问','是否存为默认值？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
        if reply == QMessageBox.Yes:
            DS1_History={}
            DS1_History['Case_Presenter']=self.lineEdit_Presenter.text()
            DS1_History['Face_Representor']=self.lineEdit_Representor.text()
            DS1_History['Face_Otherperson']=self.lineEdit_Otherperson.text()
            DS1_History['Case_Adress']=self.lineEdit_Adress.text()
            self.History['DeSetone'] = DS1_History
            with open ('History.pickle', 'wb') as History_file :
                dump(self.History, History_file)
    @pyqtSlot()
    def on_pushButton_Fordefault1_read_clicked(self):
        """
        Slot documentation goes here.
        """
        reply = QMessageBox.question(self,'询问','是否读取默认值？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
        if reply == QMessageBox.Yes:
            DS1_History=self.History['DeSetone']
            self.lineEdit_Presenter.setText(DS1_History['Case_Presenter'])
            self.lineEdit_Representor.setText(DS1_History['Face_Representor'])
            self.lineEdit_Otherperson.setText(DS1_History['Face_Otherperson'])
            self.lineEdit_Adress.setText(DS1_History['Case_Adress'])
    @pyqtSlot()
    def on_pushButton_Fordefault2_clicked(self):
        """
        Slot documentation goes here.
        """
        reply = QMessageBox.question(self,'询问','是否存为默认值？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
        if reply == QMessageBox.Yes:
            DS2_History={}
            DS2_History['Case_NumGZ']=self.lineEdit_numdocGZ.text()
            DS2_History['Case_lineNUM']=self.lineEdit_linenumber.text()
            DS2_History['Case_linePerple']=self.lineEdit_linePerple.text()
            DS2_History['Case_lineadrss']=self.lineEdit_lineadrss.text()
            DS2_History['Case_linecode']=self.lineEdit_linecode.text()
            self.History['DeSettwo'] = DS2_History
            with open ('History.pickle', 'wb') as History_file :
                dump(self.History, History_file)
    @pyqtSlot()
    def on_pushButton_Fordefault2_read_clicked(self):
        """
        Slot documentation goes here.
        """
        reply = QMessageBox.question(self,'询问','是否读取默认值？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
        if reply == QMessageBox.Yes:
            DS2_History=self.History['DeSettwo']
            self.lineEdit_linenumber.setText(DS2_History['Case_lineNUM'])
            self.lineEdit_linePerple.setText(DS2_History['Case_linePerple'])
            self.lineEdit_lineadrss.setText(DS2_History['Case_lineadrss'])
            self.lineEdit_linecode.setText(DS2_History['Case_linecode'])
    @pyqtSlot()
    def on_pushButton_replace_clicked(self):
        """
        Slot documentation goes here.
        """
        #print("你按下了快捷键ctrl+F")
        self.tabWidge_list.setCurrentIndex(0)
        

    # 主体信息导入
    @pyqtSlot(str)
    def on_comboBox_Type_activated(self, p0):
        """
        Slot documentation goes here.

        @param p0 DESCRIPTION
        @type str
        """
        self.comboBox_ClassNAME.clear()


    @pyqtSlot(str)
    def on_comboBox_ClassNAME_activated(self, p0):
        """
        Slot documentation goes here.

        @param p0 DESCRIPTION
        @type str
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        if '#' in p0:
            fname = p0.split('#')[0]
            fnumb = int(p0.split('#')[1])
            Type = self.comboBox_Type.currentText()
            xlsfile = './Date/主体信息/'+Type
            workbook = xlrd.open_workbook(xlsfile)
            sheet1=workbook.sheet_by_index(0)
            if  '医疗机构' in Type :
                fname0 = sheet1.col_values(1)[fnumb] #名称
                fnote = sheet1.col_values(12)[fnumb] #统一代码
                fnote2 = sheet1.col_values(72)[fnumb] #许可
                ftype = sheet1.col_values(5)[fnumb] #经营类型
                fadress = sheet1.col_values(8)[fnumb]  #地址
                fown = sheet1.col_values(71)[fnumb]  #经营者
                fownnote = sheet1.col_values(34)[fnumb] #身份证
                ftel =  sheet1.col_values(52)[fnumb] #联系电话
            elif '场所' in Type or '不发证单位' in Type :
                fname0 = sheet1.col_values(0)[fnumb]
                fnote = sheet1.col_values(6)[fnumb]
                fnote2 = sheet1.col_values(22)[fnumb]
                ftype = sheet1.col_values(5)[fnumb]
                fadress = sheet1.col_values(2)[fnumb]
                fown = sheet1.col_values(7)[fnumb]
                fownnote = sheet1.col_values(8)[fnumb]
                ftel =  sheet1.col_values(34)[fnumb]
            elif '学校' in Type :
                fname0 = sheet1.col_values(0)[fnumb]
                fnote = sheet1.col_values(7)[fnumb]
                #fnote2 = sheet1.col_values(22)[fnumb]
                ftype = sheet1.col_values(8)[fnumb]
                fadress = sheet1.col_values(1)[fnumb]
                fown = sheet1.col_values(9)[fnumb]
                fownnote = sheet1.col_values(11)[fnumb]
                ftel =  sheet1.col_values(12)[fnumb]
            elif '主体' in Type :
                fname0 = sheet1.col_values(1)[fnumb] #名称
                fnote = sheet1.col_values(2)[fnumb] #统一代码
                ftype = sheet1.col_values(5)[fnumb] #经营类型
                fadress = sheet1.col_values(3)[fnumb]  #地址
                fown = sheet1.col_values(7)[fnumb]  #经营者
                fownnote = ''
            if fname == fname0 :
                if len(fownnote)>16:
                    if (int(fownnote[16]) & 1) == 0:
                        fsex ='女'
                    else :
                        fsex ='男'
                else:
                    fsex= '请核对当事人身份证号码！'
                if ftype == '个体' or ftype == '私有' :
                   ftype = '个体工商户'
                if '医疗机构' in Type :
                    foutput = f"{fname}，统一社会信用代码：{fnote}，医疗机构执业许可证号：{fnote2}，经营类型：{ftype}，经营地址：{fadress}，法定代表人：{fown}（{fsex}，公民身份号码为{fownnote}），联系电话：{ftel}"
                elif '学校'in Type:
                    foutput = f"{fname}，统一社会信用代码：{fnote}，经济类型：{ftype}，地址：{fadress}，法定代表人：{fown}（{fsex}，公民身份号码为{fownnote}），联系电话：{ftel}"
                elif '主体'in Type:
                    foutput = f"{fname}，统一社会信用代码：{fnote}，经济类型：{ftype}，地址：{fadress}，法定代表人：{fown}（{fsex}"
                else:
                    foutput = f"{fname}，统一社会信用代码：{fnote}，卫生许可证号：{fnote2}，经营类型：{ftype}，经营场所：{fadress}，经营者：{fown}（{fsex}，公民身份号码为{fownnote}），联系电话：{ftel}"
                #print(foutput)
                reply = QMessageBox.question(self,'询问','是否覆盖当事人栏数据？该操作无法撤销！', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
                if reply == QMessageBox.Yes:
                    self.textEdit_litigant.setText(foutput)
            else :
                print(fname, fnumb)
        #读取数据，类型，然后读取对应xls内容，从身份证信息读取生日、性别

    @pyqtSlot(str)
    def on_comboBox_ClassNAME_editTextChanged(self, p0):
        """
        Slot documentation goes here.

        @param p0 DESCRIPTION
        @type str
        """
        # TODO: not implemented yet
        #读取文件名
        if len(p0) > 1 and self.comboBox_ClassNAME.findText(p0) == -1:
            self.comboBox_ClassNAME.clear()
            Type = self.comboBox_Type.currentText()
            if Type :
                xlsfile = './Date/主体信息/'+ Type
                workbook = xlrd.open_workbook(xlsfile)
                sheet1=workbook.sheet_by_index(0)
                if '医疗机构' in Type or '主体' in Type :
                    if len(sheet1.col_values(1)) > 1 :
                        namelist = sheet1.col_values(1)[2:]
                elif'场所'in Type or '不发证单位'in Type or '学校' in Type:
                    if len(sheet1.col_values(0)) > 1 :
                        namelist = sheet1.col_values(0)[2:]
                else:
                    QMessageBox.information(self,'错误提示', '目前仅支持医疗机构、公共场所和学校查询，\n放射和计划生育请通过医疗机构基本信息表查询！')
                    namelist =[]
                    #print(namelist)
                matchlist=[]
                #读取店名、业主名、以及统一代码列表
                #print(p0)
                for i in range(0, len(namelist)) :
                    if p0 in namelist[i]:
                        matchlist.append(namelist[i]+'#'+str(i+2))
                self.comboBox_ClassNAME.addItems(matchlist)
                if matchlist == []:
                    QMessageBox.information(self,'错误提示', f'在当前信息表中未找到含有“{p0}”字段的单位，请核对！')
                    self.comboBox_ClassNAME.clearEditText()
            else :
                print(self.comboBox_Type.count())
                if self.comboBox_Type.count() < 1:
                    reply = QMessageBox.question(self,'询问','未载入主体信息文件，是否打开存放文件夹查看？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
                    if reply == QMessageBox.Yes:
                        path_out = getcwd()
                        path_out=path.join(path_out, 'Date\主体信息')
                        open_folder(path_out)
                else:
                    QMessageBox.information(self,'错误提示', '请先选择主体信息类型！')
            #print(xlsfile)

        elif len(p0) > 0 and len(p0) < 2 :
            QMessageBox.information(self,'错误提示', '请至少输入两个字！')
            self.comboBox_ClassNAME.clearEditText()
    #段落格式调整——————————————————————————————
    @pyqtSlot()
    def on_textEdit_show_cursorPositionChanged(self):
        self.blockFormat_initA([self.textEdit_show],False)    
    @pyqtSlot()
    def on_textEdit_informed_cursorPositionChanged(self):
        self.blockFormat_initA([self.textEdit_informed],False)
    @pyqtSlot()
    def on_textEdit_heyi2_cursorPositionChanged(self):
        self.blockFormat_initA([self.textEdit_heyi2],False)
    @pyqtSlot()
    def on_textEdit_heyi_cursorPositionChanged(self):
        self.blockFormat_initA([self.textEdit_heyi],False)
    @pyqtSlot()
    def on_textEdit_process_cursorPositionChanged(self):
        self.blockFormat_initA([self.textEdit_process],False)
    @pyqtSlot()
    def on_textEdit_absractcase_2_cursorPositionChanged(self):
        self.blockFormat_initA([self.textEdit_absractcase_2],False)
    @pyqtSlot()
    def on_textEdit_absractcase_cursorPositionChanged(self):
        self.blockFormat_initA([self.textEdit_absractcase],False)
    @pyqtSlot()
    def on_textEdit_litigant_cursorPositionChanged(self):
        self.blockFormat_initA([self.textEdit_litigant],False)

    #关闭时自动保存
    def closeEvent(self, event):
        # 关闭子窗口
        if hasattr(self, 'child_window') and self.child_window:
            self.child_window.close()
        
        print('正在关闭')
        self.SAVE_context('AutoSave.pickle')
        pass


##——————------——子窗口分割线————————————————————————————————————————
class ChildWindow(QDialog, Ui_Dialog):
    """
    Class documentation goes here.
    """
    #def __init__(self, signal_manager=None,context = None, parent=None):
    def __init__(self, signal_manager=None, parent=None):
        """
        Constructor
        @param parent reference to the parent widget
        @type QWidget
        """
        super(ChildWindow, self).__init__()
        self.setupUi(self)
        self.setFixedSize(self.width(), self.height())   ##固定大小
        #self.raise_()
        self.fileslist=[]
        self.setWindowFlags(Qt.WindowMinimizeButtonHint |   # 使能最小化按钮
                            Qt.WindowCloseButtonHint )      # 使能关闭按钮)
        self.move(10, 30)
        self.parent_window = parent 
        #self.context = context if context is not None else {}
        self.context = {}
        self.signal_manager = signal_manager  # 接收 signal_manager
        #self.signal_manager.signal.context_updated.connect(self.on_context_updated)
        
    # def on_context_updated(self, context: dict):
    #     """响应上下文更新事件"""
    #     self.context = context
    #     self.plainTextEdit_Output.appendPlainText("上下文数据已更新")
        
    #docxw文件内容替换和计数
    def doc_replace(self, file, data, findlist, replacelist):
        '''
        此函数用于docx文件的批量替换
        file：文件完整路径
        date ：
        findlist：查找目标列表
        replacelist:替换内容列表
        '''
        import win32com.client
        IsWord = True
        try:
            Xapp = win32com.client.Dispatch('Word.Application')
            Xapp.Visible = 0 #后台运行
            Xapp.DisplayAlerts = 0 #不报错
        except:
            IsWord = False
        if  IsWord:
            print('使用word替换')
            wdStory = 6
            doc = Xapp.Documents.Open(file)
            for (old,new)  in zip(findlist,replacelist) :
                Xapp.Selection.EscapeKey()
                Xapp.Selection.HomeKey(Unit=wdStory)#返回到第一行
                replacetimes = 0
                while Xapp.Selection.Find.Execute(old) :
                    replacetimes += 1
                if self.checkBox_ForOnce.isChecked() :
                    Xapp.Selection.Find.Execute(old, False, False, False, False, False, True, 1, True, new, 1)#只替换一次
                    if replacetimes > 1:
                        replacetimes = 1
                else:
                    Xapp.Selection.Find.Execute(old, False, False, False, False, False, True, 1, True, new, 2)
                self.plainTextEdit_Output.appendPlainText("“{}”共替换{}次；".format(old,replacetimes ))
            if self.checkBox_ForReplace.isChecked():
                doc.Save()
            elif self.checkBox_ForReplace.isChecked() == False:
                doc.SaveAs("{}/替换结果/{}".format(data, file.split("/")[-1]))
            doc.Close()
        else:
            print('使用python替换')
            doc = docx.Document(file)
            print(doc)
            if doc and findlist :
                for (findword, replaceword) in zip(findlist, replacelist):
                    replacetimes = 0
                    for para in doc.paragraphs:
                        if self.checkBox_ForOnce.isChecked() and replacetimes ==1:
                            break
                        maxnum = para.text.count(findword)
                        for run in para.runs:
                            if self.checkBox_ForOnce.isChecked() and replacetimes ==1:
                                break
                            i = run.text.count(findword)
                            num = 1 if self.checkBox_ForOnce.isChecked() and i >1 else i
                            if run.text !="" and num > 0 :
                                run.text = run.text.replace(findword, replaceword, num)
                                replacetimes += num
                                maxnum =maxnum - num
                        if maxnum > 0 :#防止run解析错误
                            i = para.text.count(findword)
                            num = 1 if self.checkBox_ForOnce.isChecked() and i >1 else i
                            para.text = para.text.replace(findword, replaceword,num)
                            replacetimes += num

                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                i = cell.text.count(findword)
                                num = 1 if self.checkBox_ForOnce.isChecked() and i >1 else i
                                if cell.text !="" and num > 0 :
                                    cell.text = cell.text.replace(findword, replaceword, num)
                                    replacetimes += num
                    self.plainTextEdit_Output.appendPlainText("“{}”共替换{}次；".format(findword,replacetimes ))
            if self.checkBox_ForReplace.isChecked():
                doc.save("{}/{}".format(data, file.split("/")[-1]))
            elif self.checkBox_ForReplace.isChecked() == False:
                doc.save("{}/替换结果/{}".format(data, file.split("/")[-1]))

    #计数并替换函数
    def ConutReplaceText (self, findword, replaceword):
        '''
        此函数用于界面文件的批量替换并计数
        findword：查找目标
        replaceword:替换内容
        '''
        keytextlist=['Case_name','Case_abstract', 'Face_process','Face_evidence0','Face_evidencedata0','Face_evidenceGZ0','Face_evidencedataGZ0',
        'Face_evidence1','Face_evidencedata1','Face_evidenceGZ1','Face_evidencedataGZ1',
        'Face_evidence2','Face_evidencedata2','Face_evidenceGZ2','Face_evidencedataGZ2',
        'Face_evidence3','Face_evidencedata3','Face_evidenceGZ3','Face_evidencedataGZ3',
        'Face_evidence4','Face_evidencedata4','Face_evidenceGZ4','Face_evidencedataGZ4',
        'Face_evidence5','Face_evidencedata5','Face_evidenceGZ5','Face_evidencedataGZ5',
        'Face_evidence6','Face_evidencedata6','Face_evidenceGZ6','Face_evidencedataGZ6',
        'Face_evidence7','Face_evidencedata7','Face_evidenceGZ7','Face_evidencedataGZ7',
        'Face_illegfactA0', 'Face_illegfactB0', 'Face_illegfactC0', 'Face_illegfactGZA0', 'Face_illegfactGZB0','Face_illegfactGZC0',
        'Face_illegfactA1', 'Face_illegfactB1', 'Face_illegfactC1', 'Face_illegfactGZA1', 'Face_illegfactGZB1','Face_illegfactGZC1',
        'Face_illegfactA2', 'Face_illegfactB2', 'Face_illegfactC2', 'Face_illegfactGZA2', 'Face_illegfactGZB2','Face_illegfactGZC2',
        'Face_illegfactA3', 'Face_illegfactB3', 'Face_illegfactC3', 'Face_illegfactGZA3', 'Face_illegfactGZB3','Face_illegfactGZC3',
        'Face_illegfactA4', 'Face_illegfactB4', 'Face_illegfactC4', 'Face_illegfactGZA4', 'Face_illegfactGZB4','Face_illegfactGZC4',
        'Face_Punishment','Case_Presenter','Face_Representor','Face_Otherperson','Case_Adress','Case_Heyitime','Face_heyi','Face_heyi2',
        'Face_refers','Case_NumGZ','Case_lineNUM','Case_linePerple','Case_lineadrss','Case_linecode','Case_abstractGZ',
        'Case_PunishmentGZ','Case_NumJD','Face_PunishmentJD','Case_Informed','Case_ExecuWay','Case_ExecuMode','Face_show']
        keydatelist=['Face_time','Face_timeR','Face_GZtime','Face_timeEX','Face_timeJD']
        replacetimes = 0
        isreplaced = "成功"
        for key in keydatelist :
            date = self.context[key].toString(Qt.DefaultLocaleLongDate)#界面数据换成文本数据
            if self.checkBox_ForOnce.isChecked() and replacetimes ==1:
                break
            if findword == date :
                try:
                    newdate = datetime.strptime(replaceword, "%Y年%m月%d日")
                    self.context[key] =QDate(newdate)
                    replacetimes += 1
                except:
                    QMessageBox.information(self,'错误提示', f'{replaceword}日期错误！')
        for key in keytextlist:
            if self.checkBox_ForOnce.isChecked() and replacetimes ==1:
                break
            if key in self.context.keys():
                value = self.context[key]
                if value != "":
                    i = value.count(findword)
                    num = 1 if self.checkBox_ForOnce.isChecked() and i >1 else i
                    if num > 0:
                        try:
                            value = value.replace(findword, replaceword,num)
                            self.context[key] = value
                            replacetimes += num
                        except:
                            isreplaced = "失败"
                            continue
            else:
                continue
        if replacetimes == 0 and isreplaced != "失败":
            isreplaced = "无"
        elif replacetimes > 0 and isreplaced != "失败":
             isreplaced = "成功"+str(replacetimes)+"次"
        return (findword, isreplaced)
    @pyqtSlot()
    def on_pushButton_ReplaceForA_clicked(self):
        """
        Slot documentation goes here.
        """
        #文本处理
        #print(self.checkBox_ForReplace.isChecked())
        findstr = self.plainTextEdit_FindList.toPlainText().replace(';', '；')
        replacestr = self.plainTextEdit_ReplaceList.toPlainText().replace(';', '；')
        findlist = findstr.split('；')
        while '' in findlist:#去除查找列表内的空值
            findlist.remove('')
        if len(replacestr)>1 :
            replacestr = replacestr[:-1] if replacestr[-1] == '；' else  replacestr
        replacelist =[''] if replacestr == '；' else replacestr.split('；')
        self.label_find.setText(f"数量：{len(findlist)}")
        self.label_replace.setText(f"数量：{len(replacelist)}")
        if len(findlist) == len(replacelist) and findlist != ['']:
            reply = QMessageBox.Yes
            for i in range(0, len(replacelist)) :
                if replacelist[i] == '' :
                    reply = QMessageBox.question(self,'询问',f'是否将{findlist[i]}替换为空值，即删除{findlist[i]}？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
            if self.fileslist == [] and (self.radioButton_onlydoc.isChecked() or self.radioButton_both.isChecked()):
                    QMessageBox.information(self,'错误提示', '你没有选择需要替换的docx文档，请点击…按键选择！')
                    reply = QMessageBox.No
            if reply == QMessageBox.Yes:
                self.plainTextEdit_Output.clear()
                #----------------开始界面替换
                if self.radioButton_onlyui.isChecked() or self.radioButton_both.isChecked():
                    if self.parent_window and hasattr(self.parent_window, 'context'):
                        self.signal_manager.signal.context_save.emit()#发送信号，保存
                        self.context = self.parent_window.context  # 获取主窗口的 context
                        self.plainTextEdit_Output.appendPlainText("---------------------------\n程序界面全文替换开始:")
                        outputlist = map(self.ConutReplaceText, findlist, replacelist)
                        outputlist = list(outputlist)
                        for i in outputlist:
                            self.plainTextEdit_Output.appendPlainText (f"“{i[0]}”共替换{i[1]};\n")
                        self.plainTextEdit_Output.appendPlainText("---------------------------")
                        with open ('AutoSave.pickle','wb') as context_file:#保存context至缓存文件
                            dump(self.context, context_file)
                        self.signal_manager.signal.bridge.emit()#发送信号，从缓存文件读取数据到界面
                #-----------------开始文档替换
                if self.radioButton_onlydoc.isChecked() or self.radioButton_both.isChecked() :
                    #print('开始文档替换')
                    self.plainTextEdit_Output.appendPlainText("---------------------------\nWORD文档全文替换开始:")
                    indexnum=self.fileslist[0].rfind("/")
                    data = self.fileslist[0][0:indexnum]
                    if self.checkBox_ForReplace.isChecked() == False:
                        if path.exists(data+'\\替换结果') == False :
                            makedirs(data+'\\替换结果')
                    for file in self.fileslist:
                        try:
                            self.plainTextEdit_Output.appendPlainText("---------------------------\n{}开始替换：".format(file.split("/")[-1]))
                            self.doc_replace(file,data, findlist, replacelist)
                            self.plainTextEdit_Output.appendPlainText("---------------------------")
                        except:
                            QMessageBox.information(self,'错误提示', '该文档未关闭或已被删除！')
        else :
            self.plainTextEdit_Output.setPlainText("数量不相符或无替换目标，请修正后重新点击替换。")


    @pyqtSlot()
    def on_pushButton_fromExcel_clicked(self):
        """
        Slot documentation goes here.
        """
        fname = QFileDialog.getOpenFileName(self, '打开文件','./Save/',(" 存档文件(*.xls*)"))
        if fname[0]:
            reply = QMessageBox.question(self,'询问','是否从该表格读取查找/替换列表？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
            if reply == QMessageBox.Yes:
                workbook = xlrd.open_workbook(fname[0])
                sheet1=workbook.sheet_by_index(0)
                if len(sheet1.col_values(0)) > 1 and len(sheet1.col_values(1)) >1:
                    findexcellist = sheet1.col_values(0)[1:]
                    self.plainTextEdit_FindList.setPlainText('；'.join([str(x) for x in findexcellist])+'；')
                    replaceexcellist = sheet1.col_values(1)[1:]
                    self.plainTextEdit_ReplaceList.setPlainText('；'.join([str(x) for x in replaceexcellist])+'；')
                else:
                    QMessageBox.information(self,'错误提示', '你选择的表格第一列和第二列无查找/替换内容！')

    @pyqtSlot()
    def on_pushButton_SaveExcel_clicked(self):
        """
        Slot documentation goes here.
        """
        litigant = self.context['Case_name']
        list_litigant=litigant.split('，')
        litigant_name=list_litigant[0]#获取当事人名称
        reply = QMessageBox.question(self,'询问',f'此操作将会覆盖同名（{litigant_name}）表格文件？', QMessageBox.Yes | QMessageBox.No , QMessageBox.No)
        if reply == QMessageBox.Yes:
            import xlwt
            workbook = xlwt.Workbook()
            worksheet = workbook.add_sheet('sheet1', cell_overwrite_ok=True)
            #设定单元格为文本格式
            style1 = xlwt.XFStyle()  # 设置单元格格式为文本
            style1.num_format_str = '@'
            worksheet.write(0, 0,  '查找项', style1)
            worksheet.write(0, 1,  '替换项', style1)
            worksheet.write(0, 2,  '备注', style1)
            findstr = self.plainTextEdit_FindList.toPlainText().replace(';', '；')
            replacestr = self.plainTextEdit_ReplaceList.toPlainText().replace(';', '；')
            findlist = findstr.split('；')
            while '' in findlist:#去除查找列表内的空值
                findlist.remove('')
            if len(replacestr)>1 :
                replacestr = replacestr[:-1] if replacestr[-1] == '；' else  replacestr
            replacelist =[''] if replacestr == '；' else replacestr.split('；')
            if len(findlist) > 0:
                for i in range(1, len(findlist)+1):
                    worksheet.write(i, 0, findlist[i -1], style1)
            if len(replacelist) > 0:
                for i in range(1, len(replacelist)+1):
                    worksheet.write(i, 1, replacelist[i -1], style1)
            worksheet.col(0).width = 8000
            worksheet.col(1).width = 8000
            try:
                workbook.save(r"Save/%s.xls"%litigant_name)
                QMessageBox.information(self,'提示','已保存至请到Save文件夹。')
            except:
                QMessageBox.information(self,'错误提示', '请先关闭已打开的EXCEL文档')

    @pyqtSlot()
    def on_pushButton_DocLocation_clicked(self):
        """
        Slot documentation goes here.
        """

        self.fileslist, ok = QFileDialog.getOpenFileNames(self, "多文件选择", "./Output/", "Word文件 (*.docx)")
        filenamelist = []
        for file in self.fileslist:
            filenamelist.append(file.split("/")[-1].replace('.docx', ''))
        self.comboBox_DocLocation.clear()
        self.comboBox_DocLocation.addItems(filenamelist)

    @pyqtSlot()
    def on_pushButton_shift_clicked(self):
        """
        Slot documentation goes here.
        """
        findstr = self.plainTextEdit_FindList.toPlainText()
        replacestr = self.plainTextEdit_ReplaceList.toPlainText()
        if replacestr == '':
            self.plainTextEdit_ReplaceList.setPlainText(findstr)
        else :
            self.plainTextEdit_ReplaceList.setPlainText(findstr)
            self.plainTextEdit_FindList.setPlainText(replacestr)


#--------界面美化的类-------------------------
class CommonHelper:
    def __init__(self):
        pass
    @staticmethod
    def readQss(style):
        with open(style, 'r') as f:
            return f.read()

#--------主程序---------------------------------------------------------------------------
if __name__ == "__main__":
    app = QApplication(argv)
    signal_manager = SignalManager()  # 创建 SignalManager 实例
    main_window = CLASS(signal_manager)  # 将 SignalManager 传递给 CLASS
    main_window.show()
    #翻译右键菜单
    translator = QTranslator()
    translator.load("./skip/widgets_zh_CN_all.qm")
    app.installTranslator(translator)
#    #------------套用qss_____________________
    styleFile = './skip/style.qss'
    qssStyle = CommonHelper.readQss(styleFile)
    main_window.setStyleSheet(qssStyle)
    main_window.child_window.setStyleSheet(qssStyle)
    # 检查更新
    check_for_updates(VERSION)
    exit(app.exec_())
